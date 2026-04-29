from __future__ import annotations

import copy
import html
import json
import os
import re
import subprocess
import sys
import threading
import time
from datetime import datetime, date, timedelta
from pathlib import Path
from urllib import error as urlerror
from urllib import parse as urlparse
from urllib import request as urlrequest

import streamlit as st
from dotenv import load_dotenv

from core.docx_resume_export import build_docx_style_pdf_bundle, build_docx_template_pdf_bundle, pdf_backend_status
from core.resume_engine import (
    analyze_ats_score,
    generate_application_answers,
    generate_resume_content,
    improve_resume_to_target_ats,
    update_resume_content,
)
from core.storage import Storage, build_password_record, verify_password

load_dotenv()

st.set_page_config(
    page_title='TailorResume',
    page_icon='📄',
    layout='wide',
    initial_sidebar_state='collapsed',
)

APP_DIR = Path(__file__).parent
storage = Storage(APP_DIR / 'data')
APP_TITLE = 'TailorResume'
REGION_OPTIONS = ['Any', 'US', 'EU', 'LATAM']
LOW_ATS_THRESHOLD = 85
MAX_LOW_ATS_ATTEMPTS = 2


def _normalize_region(value: str) -> str:
    raw = str(value or '').strip().upper()
    if not raw or raw in {'ANY', 'ALL', 'ANYWHERE', 'GLOBAL', 'REMOTE'}:
        return 'ANY'
    return raw if raw in {'US', 'EU', 'LATAM'} else raw


def _region_label(value: str) -> str:
    normalized = _normalize_region(value)
    return 'Any' if normalized == 'ANY' else normalized


def _regions_match(profile_region: str, job_region: str) -> bool:
    profile_value = _normalize_region(profile_region)
    job_value = _normalize_region(job_region)
    return 'ANY' in {profile_value, job_value} or profile_value == job_value


def _profile_matches_job_region(profile: dict, job: dict | None) -> bool:
    if not job:
        return True
    return _regions_match(profile.get('region', ''), job.get('region', ''))


def _profile_matches_selected_job_region(profile: dict, selected_job_id: str, selected_job_region: str) -> bool:
    if not selected_job_id and _normalize_region(selected_job_region) == 'ANY':
        return True
    return _regions_match(profile.get('region', ''), selected_job_region)


def _profile_resume_upload_dir(profile_id: str) -> Path:
    return APP_DIR / 'data' / 'profile_resumes' / str(profile_id or '').strip()


def _uploaded_resume_candidate_paths(profile: dict) -> list[Path]:
    upload = profile.get('uploaded_resume') if isinstance(profile.get('uploaded_resume'), dict) else {}
    profile_id = str(profile.get('id', '') or '').strip()
    candidates: list[Path] = []

    for key in ('path', 'storage_path'):
        raw = str(upload.get(key, '') or '').strip()
        if not raw:
            continue
        path = Path(raw).expanduser()
        candidates.append(path if path.is_absolute() else APP_DIR / 'data' / path)

    relative_path = str(upload.get('relative_path', '') or '').strip()
    if relative_path:
        candidates.append(APP_DIR / 'data' / relative_path)

    filename = str(upload.get('filename', '') or '').strip()
    if profile_id and filename:
        safe_name = re.sub(r'[^A-Za-z0-9_.-]+', '_', filename).strip('._') or 'resume.docx'
        candidates.append(_profile_resume_upload_dir(profile_id) / safe_name)

    if profile_id:
        upload_dir = _profile_resume_upload_dir(profile_id)
        if upload_dir.exists():
            candidates.extend(sorted(upload_dir.glob('*.docx'), key=lambda item: item.stat().st_mtime, reverse=True))

    unique: list[Path] = []
    seen: set[str] = set()
    for candidate in candidates:
        key = str(candidate)
        if key and key not in seen:
            seen.add(key)
            unique.append(candidate)
    return unique


def _resolve_uploaded_resume_path(profile: dict) -> Path | None:
    for candidate in _uploaded_resume_candidate_paths(profile):
        try:
            if candidate.exists() and candidate.is_file() and candidate.suffix.lower() == '.docx':
                return candidate
        except OSError:
            continue
    return None


def _relative_data_path(path: Path) -> str:
    try:
        return str(path.resolve().relative_to((APP_DIR / 'data').resolve()))
    except Exception:
        return ''


def _resolved_uploaded_resume_record(profile: dict) -> dict:
    upload = copy.deepcopy(profile.get('uploaded_resume') if isinstance(profile.get('uploaded_resume'), dict) else {})
    path = _resolve_uploaded_resume_path(profile)
    if not path:
        return upload if upload else {}
    upload['path'] = str(path)
    relative_path = _relative_data_path(path)
    if relative_path:
        upload['relative_path'] = relative_path
    upload['filename'] = str(upload.get('filename', '') or path.name)
    try:
        upload['size_bytes'] = int(path.stat().st_size)
    except OSError:
        upload['size_bytes'] = int(upload.get('size_bytes', 0) or 0)
    upload.setdefault('content_type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    upload.setdefault('uploaded_at', '')
    upload.setdefault('extracted_text', '')
    return upload


def _profile_resume_status(profile: dict) -> str:
    return 'resume uploaded' if _resolve_uploaded_resume_path(profile) else 'no resume'


def _profile_has_uploaded_resume(profile: dict) -> bool:
    return _resolve_uploaded_resume_path(profile) is not None


def _format_profile_option(item: dict) -> str:
    name = item.get('name', 'Unnamed profile')
    return f"{name} [{_region_label(item.get('region', ''))}] - {_profile_resume_status(item)}"

def _format_job_option(item: dict) -> str:
    if not item.get('id'):
        return 'Manual entry'
    return f"{item.get('company', 'Unknown')} — {item.get('job_title', 'Untitled')} [{_region_label(item.get('region', ''))}]"



def init_state() -> None:
    defaults = {
        'current_user_id': '',
        'auth_notice': '',
        'last_resume': None,
        'last_exports': {},
        'last_template_id': '',
        'last_profile_id': '',
        'last_job_id': '',
        'last_job_company': '',
        'last_job_link': '',
        'last_job_description': '',
        'last_job_region': 'ANY',
        'last_target_role': '',
        'last_custom_prompt': '',
        'last_bold_keywords': '',
        'last_auto_bold_fit_keywords': False,
        'last_update_prompt': '',
        'editor_loaded_signature': '',
        'editor_pending_resume': None,
        'editor_notice': '',
        'ats_analysis_cache': {},
        'application_answers_cache': {},
        'last_ats_improve_history': [],
        'last_generator_mode': '',
        'pending_saved_resume': None,
        'saved_resume_notice': '',
        'latest_saved_resume_id': '',
        'company_message_dialog_reset_needed': False,
        'latest_saved_folder': '',
        'copy_folder_notice_path': '',
        'copy_folder_notice_pending': False,
        'job_link_fetch_error': '',
        'job_link_fetch_notice': '',
        'last_scraped_job_link': '',
        'pending_nav_page': '',
        'job_list_notice': '',
        'pending_dashboard_approved_job_id': '',
        'auth_token_value': '',
        'low_ats_attempts_by_job': {},
        'report_job_dialog_open': False,
        'report_job_dialog_target_id': '',
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


AUTH_QUERY_KEY = 'auth_token'


def _query_param_value(name: str) -> str:
    try:
        value = st.query_params.get(name, '')
    except Exception:
        try:
            value = st.experimental_get_query_params().get(name, [''])
        except Exception:
            value = ''
    if isinstance(value, (list, tuple)):
        value = value[0] if value else ''
    return str(value or '').strip()


def _set_query_param_value(name: str, value: str) -> None:
    clean = str(value or '').strip()
    try:
        if clean:
            st.query_params[name] = clean
        elif name in st.query_params:
            del st.query_params[name]
        return
    except Exception:
        pass
    try:
        params = dict(st.experimental_get_query_params())
        if clean:
            params[name] = clean
        else:
            params.pop(name, None)
        st.experimental_set_query_params(**params)
    except Exception:
        pass


def _get_auth_query_token() -> str:
    return _query_param_value(AUTH_QUERY_KEY)


def _persist_login_token(raw_token: str) -> None:
    st.session_state['auth_token_value'] = str(raw_token or '').strip()
    _set_query_param_value(AUTH_QUERY_KEY, st.session_state['auth_token_value'])


def _clear_login_token() -> None:
    raw_token = str(st.session_state.get('auth_token_value', '') or _get_auth_query_token()).strip()
    if raw_token:
        try:
            storage.revoke_auth_token(raw_token)
        except Exception:
            pass
    st.session_state['auth_token_value'] = ''
    _set_query_param_value(AUTH_QUERY_KEY, '')


def _restore_auth_from_token() -> None:
    if st.session_state.get('current_user_id'):
        return
    raw_token = _get_auth_query_token()
    if not raw_token:
        return
    user = storage.get_user_by_auth_token(raw_token)
    if user and user.get('status') == 'approved':
        st.session_state['current_user_id'] = user.get('id', '')
        st.session_state['auth_token_value'] = raw_token
        return
    st.session_state['auth_token_value'] = ''
    _set_query_param_value(AUTH_QUERY_KEY, '')


# ---------- Job scraping ----------

def _normalize_text_block(value: str) -> str:
    text = html.unescape(str(value or ''))
    text = text.replace('\r', '')
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def _strip_html_to_text(value: str) -> str:
    text = str(value or '')
    text = re.sub(r'(?is)<script[^>]*>.*?</script>', ' ', text)
    text = re.sub(r'(?is)<style[^>]*>.*?</style>', ' ', text)
    text = re.sub(r'(?i)<br\s*/?>', '\n', text)
    text = re.sub(r'(?i)</p\s*>', '\n\n', text)
    text = re.sub(r'(?i)</div\s*>', '\n', text)
    text = re.sub(r'(?i)</li\s*>', '\n', text)
    text = re.sub(r'(?i)<li[^>]*>', '• ', text)
    text = re.sub(r'(?is)<[^>]+>', ' ', text)
    text = html.unescape(text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r' *\n *', '\n', text)
    return text.strip()


def _extract_meta_content(page_html: str, attr_name: str, attr_value: str) -> str:
    pattern = rf'(?is)<meta[^>]+{attr_name}=["\']{re.escape(attr_value)}["\'][^>]+content=["\'](.*?)["\']'
    match = re.search(pattern, page_html)
    if match:
        return html.unescape(match.group(1)).strip()
    pattern_alt = rf'(?is)<meta[^>]+content=["\'](.*?)["\'][^>]+{attr_name}=["\']{re.escape(attr_value)}["\']'
    match_alt = re.search(pattern_alt, page_html)
    if match_alt:
        return html.unescape(match_alt.group(1)).strip()
    return ''


def _extract_title_from_html(page_html: str) -> str:
    candidates = [
        _extract_meta_content(page_html, 'property', 'og:title'),
        _extract_meta_content(page_html, 'name', 'twitter:title'),
        _extract_meta_content(page_html, 'name', 'title'),
    ]
    for candidate in candidates:
        if candidate:
            return candidate
    match = re.search(r'(?is)<title[^>]*>(.*?)</title>', page_html)
    if match:
        return _normalize_text_block(_strip_html_to_text(match.group(1)))
    return ''


def _iter_jobposting_nodes(node):
    if isinstance(node, dict):
        node_type = node.get('@type')
        if node_type == 'JobPosting' or (isinstance(node_type, list) and 'JobPosting' in node_type):
            yield node
        for value in node.values():
            yield from _iter_jobposting_nodes(value)
    elif isinstance(node, list):
        for item in node:
            yield from _iter_jobposting_nodes(item)


def _extract_jobposting_json_ld(page_html: str) -> dict:
    blocks = re.findall(r'(?is)<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>', page_html)
    for raw_block in blocks:
        raw_block = raw_block.strip()
        if not raw_block:
            continue
        try:
            data = json.loads(raw_block)
        except Exception:
            continue
        for job_node in _iter_jobposting_nodes(data):
            title = _normalize_text_block(job_node.get('title', ''))
            company = ''
            hiring = job_node.get('hiringOrganization') or {}
            if isinstance(hiring, dict):
                company = _normalize_text_block(hiring.get('name', ''))
            parts = []
            for key in ['description', 'responsibilities', 'qualifications', 'skills', 'experienceRequirements']:
                value = job_node.get(key, '')
                if isinstance(value, list):
                    value = '\n'.join(str(v) for v in value if v)
                parts.append(_strip_html_to_text(str(value or '')))
            description = _normalize_text_block('\n\n'.join(part for part in parts if part))
            if title or description or company:
                return {
                    'job_title': title,
                    'target_role': title,
                    'company': company,
                    'job_description': description,
                }
    return {}


def _extract_body_text(page_html: str) -> str:
    body_match = re.search(r'(?is)<body[^>]*>(.*?)</body>', page_html)
    body_html = body_match.group(1) if body_match else page_html
    text = _strip_html_to_text(body_html)
    return _normalize_text_block(text)


def scrape_job_posting(url: str) -> dict:
    cleaned_url = str(url or '').strip()
    if not cleaned_url:
        raise ValueError('Paste a job link first.')
    parsed = urlparse.urlparse(cleaned_url)
    if parsed.scheme not in {'http', 'https'} or not parsed.netloc:
        raise ValueError('Enter a valid job URL that starts with http:// or https://')

    req = urlrequest.Request(
        cleaned_url,
        headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0 Safari/537.36',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        },
    )
    try:
        with urlrequest.urlopen(req, timeout=18) as response:
            raw = response.read()
            content_type = response.headers.get_content_charset() or 'utf-8'
    except urlerror.HTTPError as exc:
        raise RuntimeError(f'Could not open the job link (HTTP {exc.code}). Some job boards block automated fetches.') from exc
    except urlerror.URLError as exc:
        raise RuntimeError('Could not reach the job link. Check the URL and your network connection.') from exc

    page_html = raw.decode(content_type, errors='ignore')
    if not page_html.strip():
        raise RuntimeError('The page returned empty content.')

    jobposting = _extract_jobposting_json_ld(page_html)
    title = jobposting.get('job_title', '')
    company = jobposting.get('company', '')
    description = jobposting.get('job_description', '')

    if not title:
        title = _extract_title_from_html(page_html)
        title = re.sub(r'\s*[|•\-–—]\s*.*$', '', title).strip()
    if not description:
        meta_description = _extract_meta_content(page_html, 'name', 'description') or _extract_meta_content(page_html, 'property', 'og:description')
        body_text = _extract_body_text(page_html)
        description = _normalize_text_block('\n\n'.join(part for part in [meta_description, body_text] if part))
    if not company:
        og_site_name = _extract_meta_content(page_html, 'property', 'og:site_name')
        company = og_site_name or parsed.netloc.replace('www.', '')

    if len(description) > 14000:
        description = description[:14000].rsplit(' ', 1)[0].strip()
    if len(description) < 150:
        raise RuntimeError('The page did not expose enough job-description text. This can happen on protected job boards like LinkedIn.')

    return {
        'job_link': cleaned_url,
        'company': company,
        'job_title': title,
        'target_role': title,
        'job_description': description,
    }


def _fetch_job_link_into_state(force: bool = False) -> None:
    raw_url = str(st.session_state.get('job_link_input', '') or '').strip()
    st.session_state['last_job_link'] = raw_url
    st.session_state['job_link_fetch_error'] = ''
    st.session_state['job_link_fetch_notice'] = ''
    if not raw_url:
        return
    if (not force) and raw_url == st.session_state.get('last_scraped_job_link') and st.session_state.get('last_job_description', '').strip():
        return
    try:
        scraped = scrape_job_posting(raw_url)
    except Exception as exc:
        st.session_state['job_link_fetch_error'] = str(exc)
        return
    st.session_state['last_scraped_job_link'] = raw_url
    st.session_state['last_job_link'] = scraped.get('job_link', raw_url)
    st.session_state['last_target_role'] = scraped.get('target_role', '')
    st.session_state['last_job_description'] = scraped.get('job_description', '')
    st.session_state['last_job_company'] = scraped.get('company', '')
    st.session_state['last_job_id'] = ''
    title = scraped.get('job_title', '').strip() or 'job posting'
    st.session_state['job_link_fetch_notice'] = f'Filled the role and job description from: {title}'


@st.cache_resource
def start_job_scrape_worker(data_dir: str) -> dict:
    worker_storage = Storage(Path(data_dir))
    stop_flag = threading.Event()

    def _loop() -> None:
        while not stop_flag.is_set():
            try:
                job = worker_storage.claim_next_pending_job_for_scrape()
                if not job:
                    time.sleep(3)
                    continue
                try:
                    scraped = scrape_job_posting(job.get('link', ''))
                    patch = {
                        'company': scraped.get('company') or job.get('company', ''),
                        'job_title': scraped.get('job_title') or job.get('job_title', ''),
                        'description': scraped.get('job_description') or job.get('description', ''),
                        'link': scraped.get('job_link') or job.get('link', ''),
                        'scrape_status': 'done',
                        'scrape_error': '',
                        'scraped_at': datetime.utcnow().isoformat() + 'Z',
                    }
                except Exception as exc:
                    patch = {
                        'scrape_status': 'error',
                        'scrape_error': str(exc),
                        'scraped_at': datetime.utcnow().isoformat() + 'Z',
                    }
                worker_storage.complete_job_scrape(job.get('id', ''), patch)
            except Exception:
                time.sleep(5)

    thread = threading.Thread(target=_loop, daemon=True, name='job-scrape-worker')
    thread.start()
    return {'started': True}


# ---------- Auth ----------

def get_current_user() -> dict | None:
    user_id = st.session_state.get('current_user_id', '')
    if not user_id:
        return None
    user = storage.get_user_by_id(user_id)
    if not user or user.get('status') != 'approved':
        st.session_state['current_user_id'] = ''
        return None
    return user


def is_admin(user: dict | None) -> bool:
    return bool(user and user.get('is_admin'))


def get_accessible_profiles(user: dict) -> list[dict]:
    profiles = storage.get_profiles()
    if is_admin(user):
        return profiles
    allowed = set(user.get('assigned_profile_ids', []))
    return [item for item in profiles if item.get('id') in allowed]


def _inject_auth_styles() -> None:
    st.markdown(
        """
        <style>
        .auth-shell {max-width: 760px; margin: 0 auto 1.25rem auto;}
        .auth-hero {text-align: center; margin: 0.35rem 0 1.1rem 0;}
        .auth-title {font-size: 3rem; font-weight: 800; letter-spacing: -0.02em; margin: 0;}
        .auth-subtitle {color: #94a3b8; margin-top: 0.35rem; font-size: 1rem;}
        .auth-card {background: rgba(15, 23, 42, 0.38); border: 1px solid rgba(148, 163, 184, 0.18); border-radius: 18px; padding: 1rem 1rem 0.35rem 1rem; box-shadow: 0 10px 30px rgba(0,0,0,0.18);}
        .auth-small {color: #94a3b8; font-size: 0.92rem;}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _get_user_by_login_identifier(identifier: str) -> dict | None:
    needle = str(identifier or '').strip().lower()
    if not needle:
        return None
    direct = storage.get_user_by_username(needle)
    if direct:
        return direct
    for user in storage.get_users():
        if str(user.get('email', '')).strip().lower() == needle:
            return user
    return None


def _password_policy_error(password: str) -> str:
    value = str(password or '')
    if len(value) < 10:
        return 'Use at least 10 characters.'
    if not re.search(r'[A-Za-z]', value) or not re.search(r'\d', value):
        return 'Include at least one letter and one number.'
    return ''


def login_screen() -> None:
    _inject_auth_styles()
    notice = st.session_state.pop('auth_notice', '')

    outer_left, outer_center, outer_right = st.columns([1.1, 2.2, 1.1])
    with outer_center:
        st.markdown(
            f"<div class='auth-shell'><div class='auth-hero'><div class='auth-title'>{APP_TITLE}</div><div class='auth-subtitle'>Secure sign in</div></div>",
            unsafe_allow_html=True,
        )
        if notice:
            st.success(notice)
        st.markdown("<div class='auth-card'>", unsafe_allow_html=True)
        tab1, tab2 = st.tabs(['Sign in', 'Request access'])
        with tab1:
            with st.form('login_form'):
                identifier = st.text_input('Username or email', placeholder='Enter your username or work email')
                password = st.text_input('Password', type='password', placeholder='Enter your password')
                submitted = st.form_submit_button('Sign in', type='primary', use_container_width=True)
            if submitted:
                user = _get_user_by_login_identifier(identifier)
                if not user:
                    st.error('No account matches that username or email.')
                    return
                if user.get('status') != 'approved':
                    st.error('This account is pending approval or has been disabled.')
                    return
                if not verify_password(password, user.get('password_salt', ''), user.get('password_hash', '')):
                    st.error('Incorrect password.')
                    return
                st.session_state['current_user_id'] = user.get('id', '')
                remember_token = storage.issue_auth_token(user.get('id', ''), ttl_days=30)
                _persist_login_token(remember_token)
                st.session_state['auth_notice'] = f"Welcome back, {user.get('full_name') or user.get('username')}"
                st.rerun()

        with tab2:
            with st.form('request_access_form'):
                full_name = st.text_input('Full name', placeholder='Your full name')
                email = st.text_input('Work email', placeholder='name@company.com')
                username = st.text_input('Requested username', placeholder='Choose a username')
                password = st.text_input('Password', type='password', placeholder='At least 10 characters with letters and numbers')
                confirm_password = st.text_input('Confirm password', type='password', placeholder='Re-enter your password')
                submitted = st.form_submit_button('Request access', type='primary', use_container_width=True)
            if submitted:
                if not full_name.strip() or not email.strip() or not username.strip() or not password.strip():
                    st.error('Full name, work email, username, and password are required.')
                    return
                if not re.match(r'^[^@\s]+@[^@\s]+\.[^@\s]+$', email.strip()):
                    st.error('Enter a valid email address.')
                    return
                if password != confirm_password:
                    st.error('Passwords do not match.')
                    return
                policy_error = _password_policy_error(password)
                if policy_error:
                    st.error(policy_error)
                    return
                if storage.get_user_by_username(username):
                    st.error('That username already exists.')
                    return
                if _get_user_by_login_identifier(email):
                    st.error('That email is already in use.')
                    return
                password_fields = build_password_record(password)
                storage.upsert_user({
                    'id': storage.make_id('user'),
                    'username': username.strip().lower(),
                    'full_name': full_name.strip(),
                    'email': email.strip(),
                    'password_hash': password_fields['password_hash'],
                    'password_salt': password_fields['password_salt'],
                    'is_admin': False,
                    'status': 'pending',
                    'assigned_profile_ids': [],
                    'created_at': datetime.utcnow().isoformat() + 'Z',
                    'approved_at': '',
                    'approved_by_user_id': '',
                    'force_password_change': False,
                })
                st.success('Access request submitted. An admin will review it before sign-in is enabled.')
        st.markdown("</div></div>", unsafe_allow_html=True)


def require_auth() -> dict:
    user = get_current_user()
    if not user:
        login_screen()
        st.stop()
    return user


def queue_nav(page_name: str) -> None:
    st.session_state['pending_nav_page'] = page_name


# ---------- Shared UI helpers ----------

def show_header(user: dict) -> None:
    if user.get('force_password_change'):
        with st.container(border=True):
            st.warning('Change your password now to replace the temporary or bootstrap credential.')
            with st.form('force_password_change_form'):
                new_password = st.text_input('New password', type='password', placeholder='At least 10 characters with letters and numbers')
                confirm_password = st.text_input('Confirm new password', type='password')
                submitted = st.form_submit_button('Update password', type='primary')
            if submitted:
                policy_error = _password_policy_error(new_password)
                if policy_error:
                    st.error(policy_error)
                elif new_password != confirm_password:
                    st.error('Passwords do not match.')
                else:
                    storage.update_user(user.get('id', ''), build_password_record(new_password) | {'force_password_change': False})
                    st.success('Password updated.')
                    st.rerun()

def _resolve_output_dir(raw_value: str) -> Path:
    value = str(raw_value or 'saved_resumes').strip() or 'saved_resumes'
    path = Path(value).expanduser()
    if not path.is_absolute():
        path = APP_DIR / path
    return path


def _parse_comma_separated_list(value: str) -> list[str]:
    seen: set[str] = set()
    items: list[str] = []
    for part in str(value or '').split(','):
        clean = part.strip()
        if not clean:
            continue
        key = clean.lower()
        if key not in seen:
            seen.add(key)
            items.append(clean)
    return items


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    output: list[str] = []
    for item in items:
        clean = str(item).strip()
        if not clean:
            continue
        key = clean.lower()
        if key not in seen:
            seen.add(key)
            output.append(clean)
    return output


def _find_index_by_id(items: list[dict], target_id: str) -> int:
    if not target_id:
        return 0
    for idx, item in enumerate(items):
        if item.get('id') == target_id:
            return idx
    return 0


def _build_file_stem(profile: dict | None) -> str:
    raw_name = str((profile or {}).get('name', '')).strip() or 'candidate'
    cleaned = ''.join(ch if ch.isalnum() else '_' for ch in raw_name)
    cleaned = re.sub(r'_+', '_', cleaned).strip('_')
    return cleaned[:80] or 'candidate'


def _render_copy_folder_notice(folder_path: str, auto_copy: bool = False) -> None:
    escaped_path = json.dumps(str(folder_path))
    auto_script = f"navigator.clipboard.writeText({escaped_path}).then(() => {{ const status = document.getElementById('copy-status'); if (status) status.textContent = 'Folder path copied to clipboard.'; }}).catch(() => {{}});" if auto_copy else ''
    st.components.v1.html(
        f"""
        <div style='margin: 0 0 12px 0;'>
          <div id='copy-box' onclick='navigator.clipboard.writeText({escaped_path}).then(() => {{ const status = document.getElementById("copy-status"); if (status) status.textContent = "Folder path copied to clipboard."; }}).catch(() => {{}});' style='cursor:pointer; background:#e8f0fe; border:1px solid #b6c9ff; color:#1d4ed8; border-radius:10px; padding:12px 14px; font-family:Arial,sans-serif;'>
            <div style='font-weight:700; margin-bottom:4px;'>Saved folder</div>
            <div style='font-size:13px; word-break:break-all;'>{str(folder_path)}</div>
            <div id='copy-status' style='font-size:12px; margin-top:6px;'>Click this blue notice to copy the folder path.</div>
          </div>
        </div>
        <script>{auto_script}</script>
        """,
        height=95,
    )


def _render_copy_value_notice(title: str, value: str, helper_text: str = 'Click to copy.') -> None:
    if not str(value or '').strip():
        return
    escaped_value = json.dumps(str(value))
    st.components.v1.html(
        f"""
        <div style='margin: 4px 0 10px 0;'>
          <div onclick='navigator.clipboard.writeText({escaped_value}).then(() => {{ const el = document.getElementById("copy-inline-status"); if (el) el.textContent = "Copied to clipboard."; }}).catch(() => {{}});' style='cursor:pointer; background:rgba(59,130,246,0.10); border:1px solid rgba(96,165,250,0.45); color:#bfdbfe; border-radius:10px; padding:10px 12px; font-family:Arial,sans-serif;'>
            <div style='font-size:12px; font-weight:700; margin-bottom:4px;'>{html.escape(title)}</div>
            <div style='font-size:12px; word-break:break-all; text-decoration:underline;'>{html.escape(str(value))}</div>
            <div id='copy-inline-status' style='font-size:11px; margin-top:4px; color:#93c5fd;'>{html.escape(helper_text)}</div>
          </div>
        </div>
        """,
        height=76,
    )


def _build_applied_map_for_user(user: dict, accessible_profiles: list[dict]) -> dict[str, set[str]]:
    generated_items = storage.get_generated_resumes()
    if not is_admin(user):
        generated_items = [item for item in generated_items if item.get('created_by_user_id') == user.get('id')]
    allowed_profile_ids = {str(item.get('id', '')) for item in accessible_profiles}
    applied_map: dict[str, set[str]] = {}
    for item in generated_items:
        job_id = str(item.get('job_id', '')).strip()
        profile_id = str(item.get('profile_id', '')).strip()
        if job_id and profile_id and (not allowed_profile_ids or profile_id in allowed_profile_ids):
            applied_map.setdefault(job_id, set()).add(profile_id)
    return applied_map


def _job_has_remaining_accessible_profiles(job: dict, accessible_profiles: list[dict], applied_map: dict[str, set[str]]) -> bool:
    if not accessible_profiles:
        return False
    applied_profiles = applied_map.get(str(job.get('id', '')).strip(), set())
    for profile in accessible_profiles:
        if _profile_matches_job_region(profile, job) and str(profile.get('id', '')) not in applied_profiles:
            return True
    return False


def _open_file_default(path_value: str) -> tuple[bool, str]:
    target = str(path_value or '').strip()
    if not target:
        return False, 'No file path was provided.'
    try:
        if sys.platform.startswith('win'):
            os.startfile(target)  # type: ignore[attr-defined]
        elif sys.platform == 'darwin':
            subprocess.Popen(['open', target])
        else:
            subprocess.Popen(['xdg-open', target])
        return True, ''
    except Exception as exc:
        return False, str(exc)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
    except Exception:
        return ''
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if value:
            parts.append(value)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts).strip()


def _save_uploaded_resume_docx(profile_id: str, uploaded_file) -> dict:
    if uploaded_file is None:
        return {}
    filename = str(getattr(uploaded_file, 'name', '') or 'resume.docx').strip()
    if not filename.lower().endswith('.docx'):
        raise ValueError('Only DOCX resume upload is supported.')
    safe_name = re.sub(r'[^A-Za-z0-9_.-]+', '_', filename).strip('._') or 'resume.docx'
    upload_dir = _profile_resume_upload_dir(profile_id)
    upload_dir.mkdir(parents=True, exist_ok=True)
    target_path = upload_dir / safe_name
    data = uploaded_file.getvalue()
    target_path.write_bytes(data)
    relative_path = _relative_data_path(target_path)
    return {
        'filename': filename,
        'content_type': str(getattr(uploaded_file, 'type', '') or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
        'size_bytes': len(data),
        'path': str(target_path),
        'relative_path': relative_path,
        'uploaded_at': datetime.utcnow().isoformat() + 'Z',
        'extracted_text': _extract_docx_text(target_path),
    }


def _pdf_export_config(settings: dict) -> dict:
    order_raw = str(settings.get('pdf_backend_order', '') or '').strip()
    backend_order = [item.strip() for item in order_raw.split(',') if item.strip()]
    if not backend_order:
        backend_order = ['docx2pdf', 'word', 'libreoffice', 'wps_custom']
    return {
        'backend_order': backend_order,
        'wps_pdf_command': str(settings.get('wps_pdf_command', '') or '').strip(),
    }


def _build_uploaded_docx_pdf_exports(resume: dict, profile: dict, app_settings: dict) -> dict:
    resolved_upload = _resolved_uploaded_resume_record(profile)
    if not resolved_upload or not Path(str(resolved_upload.get('path', ''))).exists():
        raise FileNotFoundError('no resume so must upload resume')
    export_profile = copy.deepcopy(profile)
    export_profile['uploaded_resume'] = resolved_upload
    return build_docx_style_pdf_bundle(
        resume=resume,
        profile=export_profile,
        output_dir=_resolve_output_dir(app_settings.get('download_output_dir', 'saved_resumes')),
        pdf_cfg=_pdf_export_config(app_settings),
    )


def _build_uploaded_docx_template_pdf_exports(profile: dict, app_settings: dict) -> dict:
    """Create a read-only PDF preview of the selected profile's uploaded DOCX.

    This preview runs before resume generation. It must not apply generated
    content, replace placeholders, or modify the uploaded template file.
    """
    resolved_upload = _resolved_uploaded_resume_record(profile)
    path_value = str((resolved_upload or {}).get('path', '') or '').strip()
    if not resolved_upload or not path_value or not Path(path_value).exists():
        raise FileNotFoundError('no resume so must upload resume')
    export_profile = copy.deepcopy(profile)
    export_profile['uploaded_resume'] = resolved_upload
    return build_docx_template_pdf_bundle(
        profile=export_profile,
        output_dir=_resolve_output_dir(app_settings.get('download_output_dir', 'saved_resumes')),
        pdf_cfg=_pdf_export_config(app_settings),
    )


def _uploaded_resume_signature(profile: dict) -> str:
    path = _resolve_uploaded_resume_path(profile)
    if not path:
        return f"{profile.get('id', '')}|missing"
    try:
        stat = path.stat()
        return f"{profile.get('id', '')}|{path}|{stat.st_size}|{stat.st_mtime_ns}"
    except OSError:
        return f"{profile.get('id', '')}|{path}|unknown"


def _render_application_answers_tab(resume_snapshot: dict, job_description: str, target_role: str, use_ai: bool, cache_prefix: str, default_questions: str | None = None) -> None:
    questions_key = f'{cache_prefix}_questions'
    cache_key = f'{cache_prefix}|{target_role}|{resume_snapshot.get("headline", "")}'
    if questions_key not in st.session_state:
        st.session_state[questions_key] = default_questions or '\n'.join([
            'Why are you a strong fit for this role?',
            'Tell us about your most relevant experience for this position.',
            'Why do you want to work in this role?',
        ])
    st.caption('Write one question per line. Answers use the current resume draft and the current job description, and they stay short, direct, and professionally human.')
    st.text_area('Application questions', key=questions_key, height=150)
    if st.button('Generate job application answers', key=f'{cache_prefix}_generate_answers', type='primary', use_container_width=True):
        questions = [line.strip() for line in st.session_state.get(questions_key, '').splitlines() if line.strip()]
        if not questions:
            st.error('Please enter at least one question.')
        else:
            with st.spinner('Generating answers...'):
                answer_result = generate_application_answers(
                    resume=resume_snapshot,
                    job_description=job_description,
                    questions=questions,
                    target_role=target_role,
                    use_ai=use_ai,
                )
                _record_openai_usage(answer_result, 'application_answers')
            st.session_state['application_answers_cache'][cache_key] = answer_result
    answer_result = st.session_state.get('application_answers_cache', {}).get(cache_key)
    if answer_result:
        st.caption(f"Answer mode: {answer_result.get('mode', 'n/a')}")
        for answer_index, answer_item in enumerate(answer_result.get('answers', []), start=1):
            st.markdown(f"**Q{answer_index}. {answer_item.get('question', '')}**")
            st.text_area('Answer', value=answer_item.get('answer', ''), height=120, key=f'{cache_prefix}_answer_{answer_index}')
    else:
        st.info('Generate answers to see copy-ready responses for application forms.')


# ---------- Save/finalize ----------

def _render_readable_pdf_preview(pdf_bytes: bytes, fallback_html: str = '', message: str = '') -> None:
    if not pdf_bytes:
        st.error(message or 'PDF preview is unavailable because PDF export did not return a file.')
        if fallback_html:
            st.components.v1.html(fallback_html, height=420, scrolling=True)
        return

    if message:
        st.caption(f'PDF exporter: {message}')

    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        st.warning(f'PDF image preview is unavailable because PyMuPDF is not installed: {exc}')
        if fallback_html:
            st.components.v1.html(fallback_html, height=1180, scrolling=True)
        return

    try:
        doc = fitz.open(stream=pdf_bytes, filetype='pdf')
        if doc.page_count <= 0:
            st.error('PDF preview is unavailable because the generated PDF has no pages.')
            return
        st.caption(f'Read-only PDF image preview ({doc.page_count} page{"s" if doc.page_count != 1 else ""}).')
        zoom = 2.0
        matrix = fitz.Matrix(zoom, zoom)
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            st.image(pix.tobytes('png'), caption=f'Page {page_index + 1}', use_container_width=True)
        doc.close()
    except Exception as exc:
        st.warning(f'PDF image preview failed: {exc}')
        if fallback_html:
            st.components.v1.html(fallback_html, height=1180, scrolling=True)


def _render_uploaded_resume_template_preview(profile: dict, app_settings: dict, key_prefix: str) -> None:
    """Render a read-only preview for the selected profile's uploaded DOCX."""
    resolved_upload = _resolved_uploaded_resume_record(profile)
    docx_path = _resolve_uploaded_resume_path(profile)

    if not docx_path:
        st.warning('no resume so must upload resume')
        return

    signature = _uploaded_resume_signature(profile)
    sig_key = f'{key_prefix}_template_preview_signature'
    preview_key = f'{key_prefix}_template_preview_exports'
    error_key = f'{key_prefix}_template_preview_error'

    if st.session_state.get(sig_key) != signature:
        st.session_state[sig_key] = signature
        st.session_state.pop(preview_key, None)
        st.session_state.pop(error_key, None)

    filename = str(resolved_upload.get('filename', '') or docx_path.name)
    size_bytes = int(resolved_upload.get('size_bytes', 0) or 0)
    size_label = f"{size_bytes / 1024:.1f} KB" if size_bytes else "unknown size"

    info_col, action_col = st.columns([2.2, 1.0], gap='small')
    with info_col:
        st.caption(f"Uploaded DOCX: {filename} • {size_label}")
        st.caption('Preview is read-only and uses the selected profile resume template before generation.')
    with action_col:
        try:
            st.download_button(
                'Download uploaded DOCX',
                data=docx_path.read_bytes(),
                file_name=filename,
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                use_container_width=True,
                key=f'{key_prefix}_download_uploaded_docx',
            )
        except Exception as exc:
            st.warning(f'Could not prepare DOCX download: {exc}')

    if st.button('Read uploaded resume template', key=f'{key_prefix}_read_template_button', use_container_width=True):
        try:
            with st.spinner('Creating read-only template preview...'):
                st.session_state[preview_key] = _build_uploaded_docx_template_pdf_exports(profile, app_settings)
                st.session_state.pop(error_key, None)
        except Exception as exc:
            st.session_state[preview_key] = {}
            st.session_state[error_key] = str(exc)

    preview_exports = st.session_state.get(preview_key) or {}
    preview_error = st.session_state.get(error_key, '')

    if preview_error:
        st.error(f'Template PDF preview failed: {preview_error}')

    if preview_exports:
        st.caption('Read-only preview of the uploaded DOCX template. No generated content has been applied.')
        _render_readable_pdf_preview(
            pdf_bytes=preview_exports.get('pdf', b''),
            fallback_html=preview_exports.get('html', ''),
            message=preview_exports.get('pdf_message', ''),
        )

    if preview_error or (preview_exports and not preview_exports.get('pdf')):
        extracted = str(resolved_upload.get('extracted_text', '') or '').strip()
        if not extracted:
            extracted = _extract_docx_text(docx_path)
        if extracted:
            st.caption('Text fallback preview from uploaded DOCX:')
            st.text_area(
                'Uploaded resume template text',
                value=extracted,
                height=320,
                key=f'{key_prefix}_template_text_fallback',
                disabled=True,
            )
        else:
            st.info('The uploaded DOCX exists, but text extraction returned no readable text. Use Download uploaded DOCX to open it in WPS/Word.')




def _write_saved_resume_metadata(payload: dict) -> None:
    folder = Path(payload.get('saved_folder', ''))
    if not str(folder):
        return
    folder.mkdir(parents=True, exist_ok=True)
    metadata_path = folder / 'metadata.json'
    metadata_path.write_text(json.dumps(copy.deepcopy(payload), indent=2, ensure_ascii=False), encoding='utf-8')
    message = str(payload.get('company_message', '')).strip()
    if message:
        (folder / 'company_message.txt').write_text(message, encoding='utf-8')


def _compact_resume_snapshot(resume: dict) -> dict:
    source = copy.deepcopy(resume or {})
    compact: dict = {
        'name': str(source.get('name', '')).strip(),
        'headline': str(source.get('headline', '')).strip(),
        'summary': str(source.get('summary', '')).strip(),
        'fit_keywords': [str(item).strip() for item in source.get('fit_keywords', []) if str(item).strip()],
        'technical_skills': [str(item).strip() for item in source.get('technical_skills', []) if str(item).strip()],
        'grouped_skills': {},
        'work_history': [],
        'education_history': [],
    }
    grouped_skills = source.get('grouped_skills', {}) or {}
    for key, values in grouped_skills.items():
        cleaned_values = [str(item).strip() for item in values or [] if str(item).strip()]
        if cleaned_values:
            compact['grouped_skills'][str(key).strip() or 'Other Relevant'] = cleaned_values
    for item in source.get('work_history', []) or []:
        compact['work_history'].append({
            'company_name': str(item.get('company_name', '')).strip(),
            'duration': str(item.get('duration', '')).strip(),
            'location': str(item.get('location', '')).strip(),
            'role_title': str(item.get('role_title', item.get('role', ''))).strip(),
            'role_headline': str(item.get('role_headline', '')).strip(),
            'bullets': [str(bullet).strip() for bullet in item.get('bullets', []) if str(bullet).strip()],
        })
    for item in source.get('education_history', []) or []:
        compact['education_history'].append({
            'university': str(item.get('university', '')).strip(),
            'degree': str(item.get('degree', '')).strip(),
            'duration': str(item.get('duration', '')).strip(),
            'location': str(item.get('location', '')).strip(),
        })
    return compact


def _saved_resume_payload(user: dict, profile: dict, template: dict, resume: dict, ats_analysis: dict, exports: dict, app_settings: dict) -> dict:
    saved_resume_id = storage.make_id('resume')
    created_at = datetime.utcnow().isoformat() + 'Z'

    file_stem = _build_file_stem(profile)
    pdf_name = f'{file_stem}.pdf'

    payload = {
        'saved_resume_id': saved_resume_id,
        'created_at': created_at,
        'created_date': created_at[:10],
        'created_by_user_id': user.get('id', ''),
        'created_by_username': user.get('username', ''),
        'profile_id': profile.get('id'),
        'template_id': template.get('id'),
        'job_id': st.session_state.get('last_job_id', ''),
        'job_company': st.session_state.get('last_job_company', ''),
        'job_title': st.session_state.get('last_target_role', ''),
        'job_link': st.session_state.get('last_job_link', ''),
        'job_description': st.session_state.get('last_job_description', ''),
        'job_region': st.session_state.get('last_job_region', 'ANY'),
        'target_role': st.session_state.get('last_target_role', ''),
        'resume': _compact_resume_snapshot(resume),
        'ats_score': ats_analysis.get('overall_score', 0),
        'download_filename': pdf_name,
        'download_mode': 'browser',
        'company_message': '',
        'company_message_status': 'pending',
    }
    return payload


def _finalize_saved_resume(message: str, status: str) -> None:
    payload = copy.deepcopy(st.session_state.get('pending_saved_resume') or {})
    if not payload:
        return
    payload['company_message'] = str(message or '').strip()
    payload['company_message_status'] = status
    payload['company_message_updated_at'] = datetime.utcnow().isoformat() + 'Z'
    storage.save_generated_resume(payload)
    _write_saved_resume_metadata(payload)
    st.session_state['latest_saved_resume_id'] = payload.get('saved_resume_id', '')
    st.session_state['latest_saved_folder'] = ''
    st.session_state['copy_folder_notice_path'] = ''
    st.session_state['copy_folder_notice_pending'] = False
    st.session_state['saved_resume_notice'] = 'Resume saved. The PDF was downloaded to your browser.'
    st.session_state['pending_saved_resume'] = None
    st.session_state['company_message_dialog_reset_needed'] = True


def _update_saved_resume_message(item: dict, message: str) -> None:
    saved_resume_id = item.get('saved_resume_id', '')
    if not saved_resume_id:
        return
    patch = {
        'company_message': str(message or '').strip(),
        'company_message_status': 'saved',
        'company_message_updated_at': datetime.utcnow().isoformat() + 'Z',
    }
    storage.update_generated_resume(saved_resume_id, patch)
    updated = copy.deepcopy(item)
    updated.update(patch)
    _write_saved_resume_metadata(updated)


def _submit_interview_schedule(item: dict, interviewer_name: str, interview_time: str, meeting_link: str, note: str) -> None:
    saved_resume_id = item.get('saved_resume_id', '')
    if not saved_resume_id:
        return
    existing = item.get('interview_schedule', {}) if isinstance(item.get('interview_schedule', {}), dict) else {}
    schedule_payload = {
        'interviewer_name': str(interviewer_name or '').strip(),
        'interview_time': str(interview_time or '').strip(),
        'meeting_link': str(meeting_link or '').strip(),
        'note': str(note or '').strip(),
        'submitted_at': datetime.utcnow().isoformat() + 'Z',
        'review_status': 'waiting_review',
        'reviewed_at': '',
        'reviewed_by_user_id': '',
        'reviewed_by_username': '',
        'review_note': '',
    }
    storage.update_generated_resume(saved_resume_id, {'interview_schedule': schedule_payload})
    updated = copy.deepcopy(item)
    updated['interview_schedule'] = schedule_payload
    _write_saved_resume_metadata(updated)


def _review_interview_schedule(item: dict, review_status: str, admin_user: dict, review_note: str = '') -> None:
    saved_resume_id = item.get('saved_resume_id', '')
    if not saved_resume_id:
        return
    existing = item.get('interview_schedule', {}) if isinstance(item.get('interview_schedule', {}), dict) else {}
    schedule_payload = {
        'interviewer_name': str(existing.get('interviewer_name', '')).strip(),
        'interview_time': str(existing.get('interview_time', '')).strip(),
        'meeting_link': str(existing.get('meeting_link', '')).strip(),
        'note': str(existing.get('note', '')).strip(),
        'submitted_at': str(existing.get('submitted_at', '')).strip(),
        'review_status': str(review_status or 'waiting_review').strip(),
        'reviewed_at': datetime.utcnow().isoformat() + 'Z',
        'reviewed_by_user_id': admin_user.get('id', ''),
        'reviewed_by_username': admin_user.get('username', ''),
        'review_note': str(review_note or '').strip(),
    }
    storage.update_generated_resume(saved_resume_id, {'interview_schedule': schedule_payload})
    updated = copy.deepcopy(item)
    updated['interview_schedule'] = schedule_payload
    _write_saved_resume_metadata(updated)


@st.dialog('Report job')
def _report_job_dialog() -> None:
    job_id = str(st.session_state.get('report_job_dialog_target_id', '') or '').strip()
    job = storage.get_job_by_id(job_id) if job_id else None
    if not job:
        st.warning('This job is no longer available.')
        if st.button('Close', use_container_width=True, key='report_job_close_missing'):
            st.session_state['report_job_dialog_open'] = False
            st.session_state['report_job_dialog_target_id'] = ''
            st.rerun()
        return
    st.write(f"**{job.get('company', 'Unknown')} — {job.get('job_title', 'Untitled')}**")
    if job.get('link'):
        st.caption(job.get('link', ''))
    reason = st.text_area(
        'Reason for reporting',
        key='report_job_reason_value',
        height=140,
        placeholder='Example: link is broken, role description is misleading, posting is closed, duplicate, etc.',
    )
    submit_col, cancel_col = st.columns(2)
    with submit_col:
        if st.button('Submit report', type='primary', use_container_width=True, key='report_job_submit_button'):
            cleaned_reason = str(reason or '').strip()
            if not cleaned_reason:
                st.error('Enter a reason before submitting the report.')
            else:
                user = st.session_state.get('current_user_id', '')
                user_record = storage.get_user_by_id(user) if user else None
                report_payload = {
                    'reason': cleaned_reason,
                    'reported_by_user_id': user,
                    'reported_by_username': (user_record or {}).get('username', ''),
                    'reported_at': datetime.utcnow().isoformat() + 'Z',
                    'source': 'user',
                }
                storage.add_job_report(job_id, report_payload)
                _advance_to_next_dashboard_job(job_id)
                st.session_state['job_list_notice'] = 'Job reported and flagged for admin review.'
                st.session_state['saved_resume_notice'] = 'Job reported. Moved to the next job.'
                st.session_state['report_job_dialog_open'] = False
                st.session_state['report_job_dialog_target_id'] = ''
                st.session_state.pop('report_job_reason_value', None)
                st.rerun()
    with cancel_col:
        if st.button('Cancel', use_container_width=True, key='report_job_cancel_button'):
            st.session_state['report_job_dialog_open'] = False
            st.session_state['report_job_dialog_target_id'] = ''
            st.session_state.pop('report_job_reason_value', None)
            st.rerun()


def _job_recency_sort_key(job: dict) -> str:
    """Sort key that pushes the latest-entered jobs to the top.

    Uses submitted_at when available, falling back to approved_at, then to the
    job id itself so newly created jobs (whose ids embed a timestamp prefix
    via ``make_id``) still order after older ones.
    """
    return (
        str(job.get('submitted_at', '') or '').strip()
        or str(job.get('approved_at', '') or '').strip()
        or str(job.get('id', '') or '').strip()
    )


def _enforce_low_ats_rate_limit(user: dict, job_id: str, ats_score: int) -> bool:
    """Track low-ATS attempts per job and auto-flag when the limit is hit.

    Returns True when the job has been flagged and the dashboard should advance
    to the next job. Counter is stored in session_state so it resets per
    session, matching the user-visible "tried twice in a row" expectation.
    """
    job_id = str(job_id or '').strip()
    if not job_id:
        return False
    counter_map = st.session_state.setdefault('low_ats_attempts_by_job', {})
    if ats_score >= LOW_ATS_THRESHOLD:
        counter_map.pop(job_id, None)
        return False
    counter_map[job_id] = int(counter_map.get(job_id, 0) or 0) + 1
    if counter_map[job_id] < MAX_LOW_ATS_ATTEMPTS:
        return False
    auto_reason = (
        f'Auto-flagged: ATS score stayed below {LOW_ATS_THRESHOLD} '
        f'after {MAX_LOW_ATS_ATTEMPTS} resume generations '
        f'(latest score {ats_score}/100). '
        'Job is under review because the ATS score could not exceed '
        f'{LOW_ATS_THRESHOLD} after several attempts.'
    )
    storage.add_job_report(job_id, {
        'reason': auto_reason,
        'reported_by_user_id': (user or {}).get('id', ''),
        'reported_by_username': (user or {}).get('username', ''),
        'reported_at': datetime.utcnow().isoformat() + 'Z',
        'source': 'system',
    })
    counter_map.pop(job_id, None)
    _advance_to_next_dashboard_job(job_id)
    st.session_state['saved_resume_notice'] = (
        f'ATS stayed below {LOW_ATS_THRESHOLD} after {MAX_LOW_ATS_ATTEMPTS} attempts. '
        'Job auto-reported and under review. Moved to the next job.'
    )
    return True


def _advance_to_next_dashboard_job(current_job_id: str) -> None:
    user_id = st.session_state.get('current_user_id', '')
    user_record = storage.get_user_by_id(user_id) if user_id else None
    if not user_record:
        return
    accessible = get_accessible_profiles(user_record)
    applied_map = _build_applied_map_for_user(user_record, accessible)
    approved_jobs = storage.get_jobs(include_pending=False)
    available_jobs = [
        job for job in approved_jobs
        if job.get('id') != current_job_id
        and not job.get('flagged', False)
        and not job.get('admin_applied', False)
        and _job_has_remaining_accessible_profiles(job, accessible, applied_map)
    ]
    available_jobs.sort(key=_job_recency_sort_key, reverse=True)
    if available_jobs:
        st.session_state['pending_dashboard_approved_job_id'] = available_jobs[0].get('id', '')
    else:
        st.session_state['last_job_id'] = ''
        st.session_state['last_job_company'] = ''
        st.session_state['last_job_link'] = ''
        st.session_state['last_target_role'] = ''
        st.session_state['last_job_description'] = ''
        st.session_state['last_job_region'] = 'ANY'
    st.session_state['last_resume'] = None
    st.session_state['last_exports'] = {}


@st.dialog('Save application message')
def _post_download_dialog() -> None:
    payload = st.session_state.get('pending_saved_resume') or {}
    if not payload:
        st.write('Nothing pending.')
        return
    if st.session_state.pop('company_message_dialog_reset_needed', False):
        st.session_state.pop('company_message_dialog_value', None)

    st.write(f"PDF download filename: `{payload.get('download_filename', '')}`")
    st.caption('The PDF download has started in your browser. Add the company email or application message now. If you do not have it yet, choose Update later. Cancel keeps this draft out of Generated Resumes.')
    st.text_area('Company message / application email', key='company_message_dialog_value', height=180, placeholder='Paste the email or application message from the company here...')
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button('Save message', use_container_width=True):
            message = st.session_state.get('company_message_dialog_value', '').strip()
            if not message:
                st.error('Paste the company message or choose Update later / Cancel.')
            else:
                _finalize_saved_resume(message, 'saved')
                st.rerun()
    with c2:
        if st.button('Update later', use_container_width=True):
            _finalize_saved_resume('', 'update_later')
            st.rerun()
    with c3:
        if st.button('Cancel', use_container_width=True):
            st.session_state['saved_resume_notice'] = 'Files were saved locally, but this resume was not added to Generated Resumes.'
            st.session_state['pending_saved_resume'] = None
            st.session_state['company_message_dialog_reset_needed'] = True
            st.rerun()


# ---------- Dashboard ----------

def dashboard_page(user: dict) -> None:
    show_header(user)
    accessible_profiles = get_accessible_profiles(user)
    approved_jobs = storage.get_jobs(include_pending=False)
    app_settings = storage.get_app_settings()
    default_prompt = app_settings.get('default_prompt', '')
    clean_generation = bool(app_settings.get('always_clean_generation', True))
    configured_output_dir = _resolve_output_dir(app_settings.get('download_output_dir', 'saved_resumes'))

    if not accessible_profiles:
        st.warning('No accessible profiles found. Ask an admin to assign a profile to your account.')
        return

    applied_map = _build_applied_map_for_user(user, accessible_profiles)

    pending_job_id = str(st.session_state.pop('pending_dashboard_approved_job_id', '') or '').strip()
    if pending_job_id:
        pending_job = next((job for job in approved_jobs if job.get('id') == pending_job_id), None)
        if pending_job:
            st.session_state['dashboard_approved_job_select'] = pending_job_id
            st.session_state['last_job_id'] = pending_job.get('id', '')
            st.session_state['last_job_company'] = pending_job.get('company', '')
            st.session_state['last_job_link'] = pending_job.get('link', '')
            st.session_state['last_target_role'] = pending_job.get('job_title', '')
            st.session_state['last_job_description'] = pending_job.get('description', '')
            st.session_state['last_job_region'] = _normalize_region(pending_job.get('region', ''))

    available_jobs = [
        job for job in approved_jobs
        if not job.get('flagged', False)
        and not job.get('admin_applied', False)
        and _job_has_remaining_accessible_profiles(job, accessible_profiles, applied_map)
    ]
    available_jobs.sort(key=_job_recency_sort_key, reverse=True)
    approved_job_map = {job.get('id', ''): job for job in available_jobs}
    manual_job_option = {'id': '', 'company': '', 'job_title': 'Manual entry', 'description': '', 'link': '', 'region': 'ANY'}
    approved_job_id_options = [''] + [job.get('id', '') for job in available_jobs]

    if st.session_state.get('last_job_id') and st.session_state.get('last_job_id') not in approved_job_id_options:
        st.session_state['last_job_id'] = ''
        st.session_state['last_job_company'] = ''
        st.session_state['last_job_link'] = ''
        st.session_state['last_target_role'] = ''
        st.session_state['last_job_description'] = ''
        st.session_state['last_job_region'] = 'ANY'

    desired_job_id = st.session_state.get('last_job_id', '')
    if desired_job_id not in approved_job_id_options:
        desired_job_id = ''
    if st.session_state.get('dashboard_approved_job_select') not in approved_job_id_options:
        st.session_state['dashboard_approved_job_select'] = desired_job_id

    selected_job_id = desired_job_id
    current_job_region = _normalize_region(st.session_state.get('last_job_region', 'ANY'))
    selectable_profiles = [item for item in accessible_profiles if _profile_matches_selected_job_region(item, selected_job_id, current_job_region)]
    if selected_job_id:
        applied_profiles = applied_map.get(selected_job_id, set())
        selectable_profiles = [item for item in selectable_profiles if item.get('id') not in applied_profiles]

    left_col, right_col = st.columns([0.98, 1.12], gap='large')
    with left_col:
        st.subheader('Job Details')
        job_select_col, next_job_col = st.columns([5.0, 1.15], gap='small')
        with job_select_col:
            selected_job_id = st.selectbox(
                'Approved job list (optional)',
                approved_job_id_options,
                key='dashboard_approved_job_select',
                format_func=lambda job_id: _format_job_option(approved_job_map.get(job_id) or manual_job_option),
            )
            selected_job = approved_job_map.get(selected_job_id) or manual_job_option
        with next_job_col:
            st.write('')
            if st.button('Next job', use_container_width=True, key='dashboard_next_job_button'):
                if available_jobs:
                    current_job_id = st.session_state.get('last_job_id', '')
                    available_ids = [job.get('id', '') for job in available_jobs]
                    next_index = (available_ids.index(current_job_id) + 1) % len(available_ids) if current_job_id in available_ids else 0
                    st.session_state['pending_dashboard_approved_job_id'] = available_ids[next_index]
                    st.rerun()
            report_disabled = not bool(st.session_state.get('last_job_id', ''))
            if st.button(
                'Report job',
                use_container_width=True,
                key='dashboard_report_job_button',
                disabled=report_disabled,
                help='Report this job link with a reason. Reported jobs are flagged for admin review.',
            ):
                st.session_state['report_job_dialog_open'] = True
                st.session_state['report_job_dialog_target_id'] = st.session_state.get('last_job_id', '')
                st.rerun()

        selected_job_id = selected_job.get('id', '')
        if selected_job_id and selected_job_id != st.session_state.get('last_job_id'):
            st.session_state['last_job_id'] = selected_job_id
            st.session_state['last_job_company'] = selected_job.get('company', '')
            st.session_state['last_job_link'] = selected_job.get('link', '')
            st.session_state['last_target_role'] = selected_job.get('job_title', '')
            st.session_state['last_job_description'] = selected_job.get('description', '')
            st.session_state['last_job_region'] = _normalize_region(selected_job.get('region', ''))
            st.rerun()
        elif not selected_job_id and st.session_state.get('last_job_id'):
            st.session_state['last_job_id'] = ''
            st.session_state['last_job_company'] = ''
            st.session_state['last_job_link'] = ''
            st.session_state['last_target_role'] = ''
            st.session_state['last_job_description'] = ''
            st.session_state['last_job_region'] = 'ANY'
            st.rerun()

        if selected_job_id and selected_job.get('link'):
            _render_copy_value_notice('Job link', selected_job.get('link', ''), 'Click this job link box to copy the URL.')

        current_job_region = _normalize_region(selected_job.get('region', st.session_state.get('last_job_region', 'ANY')) if selected_job_id else st.session_state.get('last_job_region', 'ANY'))

        region_col, profile_col, toggle_col = st.columns([0.95, 1.35, 0.75])
        with region_col:
            selected_region_label = st.selectbox(
                'Job market',
                REGION_OPTIONS,
                index=REGION_OPTIONS.index(_region_label(current_job_region)) if _region_label(current_job_region) in REGION_OPTIONS else 0,
                disabled=bool(selected_job_id),
                help='Manual-entry jobs can target a region. Approved jobs keep the region saved on the job.',
            )
        current_job_region = _normalize_region(selected_region_label)
        st.session_state['last_job_region'] = current_job_region

        selectable_profiles = [item for item in accessible_profiles if _profile_matches_selected_job_region(item, selected_job_id, current_job_region)]
        if selected_job_id:
            applied_profiles = applied_map.get(selected_job_id, set())
            selectable_profiles = [item for item in selectable_profiles if item.get('id') not in applied_profiles]

        if not selectable_profiles:
            st.warning('No assigned profiles match this job market or the remaining not-applied slots for this job.')
            return

        with profile_col:
            profile = st.selectbox('Profile', selectable_profiles, index=_find_index_by_id(selectable_profiles, st.session_state.get('last_profile_id')), format_func=_format_profile_option)
            if not _profile_has_uploaded_resume(profile):
                st.warning('no resume so must upload resume')
        with toggle_col:
            use_ai = st.toggle('Use OpenAI', value=True, help='If OPENAI_API_KEY is missing, generation falls back automatically.')

        if is_admin(user):
            with st.expander('Uploaded resume template preview', expanded=False):
                _render_uploaded_resume_template_preview(profile, app_settings, 'dashboard_uploaded_template_preview')

        target_role = st.text_input('Target role (optional)', value=st.session_state.get('last_target_role', ''), placeholder='Leave blank to let AI infer the best role from the job description')
        job_description = st.text_area('Job description', value=st.session_state.get('last_job_description', ''), height=330, placeholder='Paste the full job description here...')
        custom_prompt = st.text_area('Custom resume prompt (optional)', value=st.session_state.get('last_custom_prompt', ''), height=110, placeholder='Example: Keep the resume sharply aligned to backend ownership, emphasize exact named tech stacks in each company bullet, and avoid generic wording.')

        create_clicked = st.button('Create tailored resume', type='primary', use_container_width=True)

        if create_clicked:
            if not job_description.strip():
                st.error('Please paste a job description first.')
                return
            if not _profile_has_uploaded_resume(profile):
                st.error('no resume so must upload resume')
                return
            with st.spinner('Generating resume content and export files...'):
                result = generate_resume_content(
                    profile=profile,
                    job_description=job_description,
                    target_role=target_role,
                    custom_prompt=custom_prompt,
                    default_prompt=default_prompt,
                    use_ai=use_ai,
                    clean_generation=clean_generation,
                )
                _record_openai_usage(result, 'generate_resume')
                resume = result['resume']
                resume['bold_keywords'] = []
                resume['auto_bold_fit_keywords'] = False
                exports = _build_uploaded_docx_pdf_exports(resume=resume, profile=profile, app_settings=app_settings)
                ats_after_generate = analyze_ats_score(resume, job_description, target_role=target_role)
            ats_score_now = int((ats_after_generate or {}).get('overall_score', 0))
            generated_job_id = str(selected_job.get('id', '') or '').strip()
            if generated_job_id and _enforce_low_ats_rate_limit(user, generated_job_id, ats_score_now):
                st.rerun()
                return
            st.session_state['last_resume'] = resume
            st.session_state['last_exports'] = exports
            st.session_state['last_template_id'] = 'uploaded_docx_style'
            st.session_state['last_profile_id'] = profile.get('id')
            st.session_state['last_job_link'] = st.session_state.get('last_job_link', '')
            st.session_state['last_job_description'] = job_description
            st.session_state['last_target_role'] = target_role
            st.session_state['last_job_region'] = current_job_region
            st.session_state['last_custom_prompt'] = custom_prompt
            st.session_state['last_bold_keywords'] = ''
            st.session_state['last_auto_bold_fit_keywords'] = False
            st.session_state['last_update_prompt'] = ''
            st.session_state['last_generator_mode'] = result['mode']
            if selected_job.get('id'):
                st.session_state['last_job_id'] = selected_job.get('id', '')
                st.session_state['last_job_company'] = selected_job.get('company', '')
            _queue_editor_reload(resume, f"Resume generated in {result['mode']} mode.")
            st.rerun()

    with right_col:
        notice = st.session_state.pop('saved_resume_notice', '')
        if notice:
            st.success(notice)

        current_ats = None
        if st.session_state.get('last_resume') and st.session_state.get('last_job_description', '').strip():
            current_ats = analyze_ats_score(
                st.session_state.get('last_resume') or {},
                st.session_state.get('last_job_description', ''),
                target_role=st.session_state.get('last_target_role', ''),
            )
        current_profile = storage.get_profile_by_id(st.session_state.get('last_profile_id')) or selectable_profiles[0]
        current_template = {'id': 'uploaded_docx_style', 'name': 'Uploaded DOCX style'}
        title_col, action_col = st.columns([4.3, 1.5])
        with title_col:
            if current_ats:
                st.subheader(f"Generated Resume • ATS {current_ats.get('overall_score', 0)}/100")
            else:
                st.subheader('Generated Resume')
        with action_col:
            can_save_pdf = bool(st.session_state.get('last_resume')) and bool(current_ats) and int(current_ats.get('overall_score', 0)) > 90 and bool((st.session_state.get('last_exports') or {}).get('pdf'))
            save_help = "Downloads the styled PDF generated from the uploaded DOCX. PDF download unlocks when ATS is over 90." if can_save_pdf else 'PDF download unlocks when ATS is over 90.'
            download_clicked = st.download_button(
                'Download PDF',
                data=(st.session_state.get('last_exports') or {}).get('pdf', b''),
                file_name=f"{_build_file_stem(current_profile)}.pdf",
                mime='application/pdf',
                use_container_width=True,
                disabled=not can_save_pdf,
                help=save_help,
                key='dashboard_download_pdf_button',
            )
            if download_clicked:
                payload = _saved_resume_payload(user, current_profile, current_template, st.session_state.get('last_resume') or {}, current_ats or {}, st.session_state.get('last_exports') or {}, app_settings)
                st.session_state['saved_resume_notice'] = 'PDF download started.'
                st.session_state['pending_saved_resume'] = payload
                st.session_state['company_message_dialog_reset_needed'] = True
                st.rerun()
        if st.session_state.get('last_resume'):
            resume = st.session_state['last_resume']
            pending_resume = st.session_state.pop('editor_pending_resume', None)
            if pending_resume is not None:
                _load_editor_from_resume(pending_resume, force=True)
            else:
                _load_editor_from_resume(resume)
            notice = st.session_state.pop('editor_notice', '')
            if notice:
                st.success(notice)
            exports = st.session_state['last_exports']
            tab_labels = ['Preview', 'Edit & Fix', 'Exports', 'ATS Notes']
            if is_admin(user):
                tab_labels.append('Job Application Answers')
            tab_labels.extend(['Structured Data', 'Source Profile'])
            tab_objects = dict(zip(tab_labels, st.tabs(tab_labels)))
            with tab_objects['Preview']:
                st.caption('Read-only PDF preview. The PDF is generated from the uploaded DOCX style; only title, summary, skills, and experience content are changed.')
                _render_readable_pdf_preview(
                    pdf_bytes=(exports or {}).get('pdf', b''),
                    fallback_html=(exports or {}).get('html', ''),
                    message=(exports or {}).get('pdf_message', ''),
                )
            with tab_objects['Edit & Fix']:
                _edit_and_fix_tab(current_profile, current_template, st.session_state.get('last_job_description', ''), st.session_state.get('last_target_role', ''), st.session_state.get('last_custom_prompt', ''), default_prompt, use_ai, clean_generation)
            with tab_objects['Exports']:
                st.markdown('**Download behavior**')
                pdf_message = (st.session_state.get('last_exports') or {}).get('pdf_message', '')
                if pdf_message:
                    st.caption(f'PDF exporter: {pdf_message}')
                if not (st.session_state.get('last_exports') or {}).get('pdf'):
                    st.error('PDF export failed. Check App Settings > PDF backend order / WPS custom PDF command.')
                st.info("In the deployed app, Download PDF sends the file to the user's browser and local Downloads flow.")
                if current_ats and int(current_ats.get('overall_score', 0)) > 90:
                    st.success('Use the Download PDF button next to the Generated Resume title to download this ATS-qualified resume.')
                else:
                    st.warning('PDF download becomes available only when the current ATS score is over 90.')
                latest_saved = st.session_state.get('latest_saved_resume_id', '')
                if latest_saved:
                    st.caption(f'Latest saved resume id: {latest_saved}')
            with tab_objects['ATS Notes']:
                _dashboard_ats_notes_tab(current_profile, current_template, resume, st.session_state.get('last_job_description', ''), st.session_state.get('last_target_role', ''), st.session_state.get('last_custom_prompt', ''), default_prompt, use_ai, clean_generation)
            if 'Job Application Answers' in tab_objects:
                with tab_objects['Job Application Answers']:
                    _render_application_answers_tab(
                        resume_snapshot=resume,
                        job_description=st.session_state.get('last_job_description', ''),
                        target_role=st.session_state.get('last_target_role', ''),
                        use_ai=use_ai,
                        cache_prefix='dashboard_current_resume_answers',
                    )
            with tab_objects['Structured Data']:
                st.json(resume)
            with tab_objects['Source Profile']:
                st.json(current_profile)
        else:
            st.info('Generate a resume from the left panel to see the preview, edit tools, ATS guidance, and export options here.')

    if st.session_state.get('pending_saved_resume'):
        _post_download_dialog()

    if st.session_state.get('report_job_dialog_open'):
        _report_job_dialog()


# ---------- Edit / fix ----------

def _edit_and_fix_tab(profile: dict, template: dict, job_description: str, target_role: str, custom_prompt: str, default_prompt: str, use_ai: bool, clean_generation: bool) -> None:
    st.caption('Edit the current draft directly, then apply manual changes or send a fix request to OpenAI using the current draft as the starting point.')
    st.markdown('#### Manual draft editor')
    st.text_input('Headline', key='editor_headline')
    st.text_area('Summary', key='editor_summary', height=150)
    c1, c2 = st.columns(2)
    with c1:
        st.text_area('Technical skills (comma separated)', key='editor_technical_skills', height=120)
    with c2:
        st.text_input('Fit keywords (comma separated)', key='editor_fit_keywords')
    st.text_area('Skill groups (one line per group: Category: item1, item2)', key='editor_skill_groups', height=150)
    work_count = int(st.session_state.get('editor_work_count', 0))
    for idx in range(work_count):
        with st.expander(f'Work Experience {idx + 1}', expanded=(idx == 0)):
            a, b, c = st.columns([1.2, 1.0, 1.0])
            with a:
                st.text_input('Company', key=f'editor_job_company_{idx}')
            with b:
                st.text_input('Duration', key=f'editor_job_duration_{idx}')
            with c:
                st.text_input('Location', key=f'editor_job_location_{idx}')
            st.text_input('Role title', key=f'editor_job_role_{idx}')
            st.text_area('Role headline', key=f'editor_job_headline_{idx}', height=80)
            st.text_area('Bullets (one per line)', key=f'editor_job_bullets_{idx}', height=180)
    st.markdown('#### Education')
    st.text_area('Education history (University | Degree | Duration | Location)', key='editor_education', height=120)
    m1, m2 = st.columns([1, 1])
    with m1:
        if st.button('Apply manual edits', use_container_width=True):
            updated_resume = _resume_from_editor(st.session_state.get('last_resume') or {})
            updated_resume['bold_keywords'] = (st.session_state.get('last_resume') or {}).get('bold_keywords', [])
            updated_resume['auto_bold_fit_keywords'] = bool((st.session_state.get('last_resume') or {}).get('auto_bold_fit_keywords', False))
            exports = _build_uploaded_docx_pdf_exports(resume=updated_resume, profile=profile, app_settings=storage.get_app_settings())
            st.session_state['last_resume'] = updated_resume
            st.session_state['last_exports'] = exports
            st.session_state['last_generator_mode'] = 'manual-edit'
            _queue_editor_reload(updated_resume, 'Manual changes applied. Preview and downloads now use the updated draft.')
            st.rerun()
    with m2:
        if st.button('Reset editor to current draft', use_container_width=True):
            _queue_editor_reload(st.session_state.get('last_resume') or {}, 'Editor reset to the current draft.')
            st.rerun()
    st.markdown('---')
    st.markdown('#### Fix with OpenAI')
    st.text_area('Fix requirement prompt', key='editor_update_prompt', height=120, placeholder='Example: Tighten the summary for a backend role, reduce generic wording, and make every bullet mention exact supported stacks where justified by the source profile.')
    if st.button('Update resume with OpenAI', type='primary', use_container_width=True):
        fix_prompt = st.session_state.get('editor_update_prompt', '').strip()
        if not fix_prompt:
            st.error('Please enter a fix requirement prompt first.')
            return
        current_draft = _resume_from_editor(st.session_state.get('last_resume') or {})
        current_draft['bold_keywords'] = (st.session_state.get('last_resume') or {}).get('bold_keywords', [])
        current_draft['auto_bold_fit_keywords'] = bool((st.session_state.get('last_resume') or {}).get('auto_bold_fit_keywords', False))
        with st.spinner('Updating current draft with OpenAI...'):
            result = update_resume_content(profile=profile, job_description=job_description, current_resume=current_draft, fix_prompt=fix_prompt, target_role=target_role, custom_prompt=custom_prompt, default_prompt=default_prompt, use_ai=use_ai, clean_generation=clean_generation)
            _record_openai_usage(result, 'update_resume')
            updated_resume = result['resume']
            updated_resume['bold_keywords'] = current_draft.get('bold_keywords', [])
            updated_resume['auto_bold_fit_keywords'] = bool(current_draft.get('auto_bold_fit_keywords', False))
            exports = _build_uploaded_docx_pdf_exports(resume=updated_resume, profile=profile, app_settings=storage.get_app_settings())
        st.session_state['last_resume'] = updated_resume
        st.session_state['last_exports'] = exports
        st.session_state['last_update_prompt'] = fix_prompt
        st.session_state['last_generator_mode'] = result['mode']
        _queue_editor_reload(updated_resume, f"Resume updated in {result['mode']} mode.")
        st.rerun()


# ---------- Admin pages ----------

def _profile_name_conflict(name: str, current_profile_id: str = '') -> dict | None:
    clean_name = re.sub(r'\s+', ' ', str(name or '')).strip().casefold()
    current_profile_id = str(current_profile_id or '').strip()
    if not clean_name:
        return None
    for profile in storage.get_profiles():
        existing_id = str(profile.get('id', '') or '').strip()
        existing_name = re.sub(r'\s+', ' ', str(profile.get('name', '') or '')).strip().casefold()
        if existing_name == clean_name and existing_id != current_profile_id:
            return profile
    return None


def profile_settings_page(user: dict) -> None:
    if not is_admin(user):
        st.error('Only admins can manage profiles.')
        return
    show_header(user)
    st.subheader('Profile Settings')
    profiles = storage.get_profiles()
    profile_labels = [f"{p.get('name', 'Unnamed')} - {_profile_resume_status(p)}" for p in profiles]
    mode = st.radio('Profile action', ['Edit existing', 'Create new'], horizontal=True)
    if mode == 'Edit existing' and profiles:
        selected_label = st.selectbox('Choose profile', profile_labels)
        selected_index = profile_labels.index(selected_label)
        selected_profile = profiles[selected_index]
    else:
        selected_profile = {'id': '', 'name': '', 'email': '', 'phone': '', 'location': '', 'region': 'ANY', 'linkedin': '', 'portfolio': '', 'summary_seed': '', 'technical_skills': [], 'work_history': [], 'education_history': [], 'uploaded_resume': {}}

    upload_info = _resolved_uploaded_resume_record(selected_profile)
    if _profile_has_uploaded_resume(selected_profile):
        st.success(f"resume uploaded: {upload_info.get('filename', 'resume.docx')}")
        with st.expander('Read uploaded resume template before generate', expanded=False):
            _render_uploaded_resume_template_preview(selected_profile, storage.get_app_settings(), 'profile_settings_template_preview')
    else:
        st.warning('no resume so must upload resume')

    with st.form('profile_form'):
        c1, c2 = st.columns(2)
        with c1:
            name = st.text_input('Full name', value=selected_profile.get('name', ''))
            email = st.text_input('Email', value=selected_profile.get('email', ''))
            phone = st.text_input('Phone', value=selected_profile.get('phone', ''))
        with c2:
            location = st.text_input('Location', value=selected_profile.get('location', ''))
            region = st.selectbox('Profile market', REGION_OPTIONS, index=REGION_OPTIONS.index(_region_label(selected_profile.get('region', 'ANY'))) if _region_label(selected_profile.get('region', 'ANY')) in REGION_OPTIONS else 0)
            linkedin = st.text_input('LinkedIn', value=selected_profile.get('linkedin', ''))
            portfolio = st.text_input('Portfolio / GitHub', value=selected_profile.get('portfolio', ''))
        uploaded_resume = st.file_uploader('Upload resume DOCX', type=['docx'], help='Required. The generated PDF keeps this DOCX style and only replaces allowed resume content sections.')
        st.caption('Only DOCX uploads are accepted. New profiles cannot be saved without a resume DOCX.')
        skills_text = st.text_area('Technical skills (comma separated, optional helper for AI)', value=', '.join(selected_profile.get('technical_skills', [])), height=90)
        summary_seed = st.text_area('About / summary seed (optional helper for AI)', value=selected_profile.get('summary_seed', ''), height=100)
        st.markdown('### Work history helper data (optional)')
        st.caption('The uploaded DOCX is the source style. These fields are optional AI helper data if you want stronger structured generation.')
        work_history_text = st.text_area('Work history', value=_serialize_work_history(selected_profile.get('work_history', [])), height=220)
        st.markdown('### Education history helper data (optional)')
        education_text = st.text_area('Education history', value=_serialize_education_history(selected_profile.get('education_history', [])), height=120)
        submitted = st.form_submit_button('Save profile', type='primary')
    if submitted:
        if not name.strip():
            st.error('Profile name is required.')
            return
        existing_profile = _profile_name_conflict(name, selected_profile.get('id', ''))
        if existing_profile:
            st.error(f"Profile name already exists: {existing_profile.get('name', '').strip() or 'Unnamed profile'}. Use a different profile name.")
            return
        profile_id = selected_profile.get('id') or storage.make_id('profile')
        uploaded_resume_record = copy.deepcopy(upload_info) if upload_info else {}
        if uploaded_resume is not None:
            try:
                uploaded_resume_record = _save_uploaded_resume_docx(profile_id, uploaded_resume)
            except Exception as exc:
                st.error(str(exc))
                return
        resolved_existing_path = Path(str(uploaded_resume_record.get('path', ''))) if uploaded_resume_record.get('path') else None
        if not uploaded_resume_record or not resolved_existing_path or not resolved_existing_path.exists():
            st.error('no resume so must upload resume')
            return
        payload = {
            'id': profile_id,
            'name': name.strip(),
            'email': email.strip(),
            'phone': phone.strip(),
            'location': location.strip(),
            'region': _normalize_region(region),
            'linkedin': linkedin.strip(),
            'portfolio': portfolio.strip(),
            'default_template_id': '',
            'summary_seed': summary_seed.strip(),
            'technical_skills': [item.strip() for item in skills_text.split(',') if item.strip()],
            'work_history': _parse_work_history(work_history_text),
            'education_history': _parse_education_history(education_text),
            'uploaded_resume': uploaded_resume_record,
        }
        storage.upsert_profile(payload)
        st.success('Profile saved.')
        st.rerun()
    if mode == 'Edit existing' and profiles:
        if st.button('Delete selected profile'):
            storage.delete_profile(selected_profile.get('id'))
            st.success('Profile deleted.')
            st.rerun()

def app_settings_page(user: dict) -> None:
    if not is_admin(user):
        st.error('Only admins can manage app settings.')
        return
    show_header(user)
    st.subheader('App Settings')
    settings = storage.get_app_settings()
    with st.form('app_settings_form'):
        default_prompt = st.text_area('Default prompt', value=settings.get('default_prompt', ''), height=220, placeholder='Default resume guidance that should apply to every generation...')
        download_output_dir = st.text_input('Server save folder (local desktop mode only)', value=settings.get('download_output_dir', 'saved_resumes'), help='Temporary DOCX/PDF files are created server-side before browser download.')
        pdf_backend_order = st.text_input('PDF backend order', value=settings.get('pdf_backend_order', 'docx2pdf, word, libreoffice, wps_custom'), help='Comma-separated. For Windows + WPS, configure wps_custom when docx2pdf/Word is unavailable.')
        wps_pdf_command = st.text_input('WPS custom PDF command', value=settings.get('wps_pdf_command', ''), help='Optional. Example: "C:\\Path\\to\\wps_export.bat" "{input}" "{output}"')
        submitted = st.form_submit_button('Save app settings', type='primary')
    if submitted:
        storage.save_app_settings({'default_prompt': default_prompt.strip(), 'always_clean_generation': True, 'download_output_dir': download_output_dir.strip() or 'saved_resumes', 'pdf_backend_order': pdf_backend_order.strip() or 'docx2pdf, word, libreoffice, wps_custom', 'wps_pdf_command': wps_pdf_command.strip()})
        st.success('App settings saved.')
        st.rerun()
    with st.expander('Detected PDF export backends'):
        for line in pdf_backend_status(_pdf_export_config(settings)):
            st.write(f'- {line}')


def _profile_assignment_owner_map(users: list[dict], exclude_user_id: str = "") -> dict[str, str]:
    owner_map: dict[str, str] = {}
    for item in users:
        if item.get('id') == exclude_user_id:
            continue
        if item.get('status') not in {'approved', 'pending'}:
            continue
        owner_name = item.get('full_name') or item.get('username') or 'another user'
        for profile_id in item.get('assigned_profile_ids', []) or []:
            if profile_id:
                owner_map[str(profile_id)] = str(owner_name)
    return owner_map


def _available_profiles_for_user_assignment(profiles: list[dict], users: list[dict], current_user_id: str, current_assigned_ids: list[str] | None = None) -> tuple[list[dict], dict[str, str]]:
    current_assigned = {str(item) for item in (current_assigned_ids or []) if str(item)}
    owner_map = _profile_assignment_owner_map(users, exclude_user_id=current_user_id)
    available: list[dict] = []
    for profile in profiles:
        profile_id = str(profile.get('id', ''))
        if not profile_id:
            continue
        if profile_id in current_assigned or profile_id not in owner_map:
            available.append(profile)
    return available, owner_map


def _assigned_profile_help_text(selected_profiles: list[dict], owner_map: dict[str, str]) -> str:
    taken = []
    for profile_id, owner in owner_map.items():
        taken.append((profile_id, owner))
    if not taken:
        return 'Each profile can be assigned to only one user.'
    return 'Each profile can be assigned to only one user. Profiles already assigned to other users are hidden from this picker.'


def _record_openai_usage(result: dict, kind: str) -> None:
    """Record one OpenAI call for the current user when the result came from a real API call."""
    mode = str((result or {}).get('mode', '')).strip().lower()
    if not mode.startswith('openai'):
        return
    user_id = str(st.session_state.get('current_user_id', '') or '').strip()
    if not user_id:
        return
    try:
        storage.record_openai_call(user_id, kind=kind)
    except Exception:
        pass


def _record_openai_usage_for_improve(result: dict) -> None:
    """Record one OpenAI call per ATS-improve round that actually hit the API."""
    user_id = str(st.session_state.get('current_user_id', '') or '').strip()
    if not user_id:
        return
    history = (result or {}).get('history', []) or []
    for entry in history:
        round_mode = str((entry or {}).get('mode', '')).strip().lower()
        if round_mode.startswith('openai'):
            try:
                storage.record_openai_call(user_id, kind='ats_improve_round')
            except Exception:
                pass


def _safe_parse_datetime(value: str) -> datetime | None:
    raw = str(value or '').strip()
    if not raw:
        return None
    try:
        return datetime.fromisoformat(raw.replace('Z', '+00:00'))
    except Exception:
        return None


def _week_start_for(dt_value: datetime) -> date:
    return dt_value.date() - timedelta(days=dt_value.weekday())


def _application_metrics_rows() -> list[dict]:
    users = storage.get_users()
    user_map = {str(item.get('id', '')): item for item in users}
    records = storage.get_generated_resumes()
    rows: list[dict] = []
    for item in records:
        created_dt = _safe_parse_datetime(item.get('created_at', ''))
        if not created_dt:
            continue
        user_id = str(item.get('created_by_user_id', '')).strip()
        user_obj = user_map.get(user_id, {})
        interview_schedule = item.get('interview_schedule', {}) if isinstance(item.get('interview_schedule', {}), dict) else {}
        submitted_dt = _safe_parse_datetime(interview_schedule.get('submitted_at', ''))
        rows.append({
            'user_id': user_id,
            'username': user_obj.get('username', '') or item.get('created_by_username', ''),
            'full_name': user_obj.get('full_name', '') or item.get('created_by_username', ''),
            'week_start': _week_start_for(created_dt),
            'created_at': created_dt,
            'job_company': item.get('job_company', ''),
            'job_title': item.get('job_title', '') or item.get('target_role', ''),
            'profile_id': item.get('profile_id', ''),
            'schedule_submitted_at': submitted_dt,
            'schedule_review_status': str(interview_schedule.get('review_status', 'not_submitted') or 'not_submitted').strip(),
        })
    return rows


def _metrics_available_week_starts(rows: list[dict]) -> list[date]:
    current_week_start = date.today() - timedelta(days=date.today().weekday())
    return sorted({row['week_start'] for row in rows if row.get('week_start')} | {current_week_start}, reverse=True)


def _week_label(week_start: date) -> str:
    week_end = week_start + timedelta(days=6)
    return f"{week_start.strftime('%Y-%m-%d')} to {week_end.strftime('%Y-%m-%d')}"


def _metrics_week_selector(rows: list[dict], key: str, label_visibility: str = 'collapsed') -> date:
    available_week_starts = _metrics_available_week_starts(rows)
    current_week_start = date.today() - timedelta(days=date.today().weekday())
    week_options = {_week_label(week_start): week_start for week_start in available_week_starts}
    week_labels = list(week_options.keys())
    default_index = available_week_starts.index(current_week_start) if current_week_start in available_week_starts else 0
    selected_week_label = st.selectbox(
        'Week',
        week_labels,
        index=default_index,
        key=key,
        label_visibility=label_visibility,
    )
    return week_options[selected_week_label]


def _application_metrics_column_config(include_openai: bool = True) -> dict:
    """Annotate each weekly metrics column with a hover description.

    Admin view (``include_openai=True``) shows ``applications/openai_calls`` per
    day. User view (``include_openai=False``) shows applications only — the
    OpenAI call counts are admin-internal data.
    """
    if include_openai:
        day_help = 'Applications saved that day / OpenAI API calls made that day for this user.'
        sum_help = 'Weekly total: applications saved / OpenAI calls.'
        day_column = lambda label: st.column_config.TextColumn(label, help=day_help)
        sum_column = st.column_config.TextColumn('Sum', help=sum_help)
    else:
        day_help = 'Applications saved that day.'
        sum_help = 'Weekly total of applications saved.'
        day_column = lambda label: st.column_config.NumberColumn(label, help=day_help)
        sum_column = st.column_config.NumberColumn('Sum', help=sum_help)
    return {
        'User': st.column_config.TextColumn('User', help='Approved user (full name or username).'),
        'Mon': day_column('Mon'),
        'Tue': day_column('Tue'),
        'Wed': day_column('Wed'),
        'Thu': day_column('Thu'),
        'Fri': day_column('Fri'),
        'Sat': day_column('Sat'),
        'Sun': day_column('Sun'),
        'Sum': sum_column,
        'Schedules': st.column_config.NumberColumn('Schedules', help='Interview schedule submissions made this week.'),
    }


def _openai_call_index() -> dict[tuple[str, date], int]:
    """Aggregate OpenAI call counts by (user_id, date)."""
    index: dict[tuple[str, date], int] = {}
    for entry in storage.get_openai_calls():
        recorded_dt = _safe_parse_datetime(entry.get('recorded_at', ''))
        user_id = str(entry.get('user_id', '') or '').strip()
        if not recorded_dt or not user_id:
            continue
        key = (user_id, recorded_dt.date())
        index[key] = index.get(key, 0) + 1
    return index


def _build_weekly_summary_rows(rows: list[dict], users: list[dict], selected_week_start: date, include_openai: bool = True) -> list[dict]:
    day_offsets = [
        ('Mon', 0),
        ('Tue', 1),
        ('Wed', 2),
        ('Thu', 3),
        ('Fri', 4),
        ('Sat', 5),
        ('Sun', 6),
    ]

    def _day_count(user_id: str, day_index: int) -> int:
        target_date = selected_week_start + timedelta(days=day_index)
        total = 0
        for row in rows:
            if row['user_id'] != user_id:
                continue
            if row['created_at'].date() == target_date:
                total += 1
        return total

    def _schedule_count(user_id: str) -> int:
        week_end = selected_week_start + timedelta(days=6)
        total = 0
        for row in rows:
            submitted_dt = row.get('schedule_submitted_at')
            if row['user_id'] != user_id or not submitted_dt:
                continue
            if selected_week_start <= submitted_dt.date() <= week_end:
                total += 1
        return total

    openai_index = _openai_call_index() if include_openai else {}

    summary_rows: list[dict] = []
    for member in users:
        uid = str(member.get('id', ''))
        member_name = member.get('full_name') or member.get('username') or 'Unknown'
        row = {'User': member_name}
        week_app_total = 0
        week_openai_total = 0
        for label, offset in day_offsets:
            day_date = selected_week_start + timedelta(days=offset)
            apps = _day_count(uid, offset)
            if include_openai:
                openai_count = int(openai_index.get((uid, day_date), 0))
                row[label] = f'{apps}/{openai_count}'
                week_openai_total += openai_count
            else:
                row[label] = apps
            week_app_total += apps
        row['Sum'] = f'{week_app_total}/{week_openai_total}' if include_openai else week_app_total
        row['Schedules'] = _schedule_count(uid)
        summary_rows.append(row)
    return sorted(summary_rows, key=lambda item: str(item.get('User', '')).lower())


def _render_application_metrics_tab() -> None:
    rows = _application_metrics_rows()
    users = [item for item in storage.get_users() if item.get('status') == 'approved']
    if not users:
        st.info('No approved users found.')
        return

    title_col, week_col = st.columns([3, 2])
    with title_col:
        st.markdown('**Applications by user**')
    with week_col:
        selected_week_start = _metrics_week_selector(rows, key='metrics_selected_week')

    summary_rows = _build_weekly_summary_rows(rows, users, selected_week_start)
    st.dataframe(
        summary_rows,
        use_container_width=True,
        hide_index=True,
        column_config=_application_metrics_column_config(),
    )


def _render_schedule_reviews_tab(admin_user: dict) -> None:
    items = storage.get_generated_resumes()
    status_options = ['waiting_review', 'useful', 'declined']
    selected_status = st.selectbox('Review status', status_options, key='schedule_review_status_filter')
    filtered_items = [
        item for item in items
        if str((item.get('interview_schedule') or {}).get('review_status', 'not_submitted')) == selected_status
    ]
    if not filtered_items:
        st.info('No interview schedule submissions matched that status.')
        return

    profiles_map = {item.get('id'): item for item in storage.get_profiles()}
    for item in sorted(filtered_items, key=lambda record: str((record.get('interview_schedule') or {}).get('submitted_at', '') or record.get('created_at', '')), reverse=True):
        schedule = item.get('interview_schedule', {}) or {}
        profile_name = (profiles_map.get(item.get('profile_id')) or {}).get('name', '') or 'Unknown profile'
        label = f"{_generated_resume_display_title(item)} • {profile_name} • {schedule.get('review_status', 'not_submitted')}"
        with st.expander(label):
            st.write(f"Company: {item.get('job_company', '') or '—'}")
            st.write(f"Job title: {item.get('job_title', '') or item.get('target_role', '') or '—'}")
            st.write(f"Created by: {item.get('created_by_username', '') or '—'}")
            st.write(f"Interviewer: {schedule.get('interviewer_name', '') or '—'}")
            st.write(f"Interview time: {schedule.get('interview_time', '') or '—'}")
            if schedule.get('meeting_link'):
                st.write(f"Meeting link: {schedule.get('meeting_link')}")
            if schedule.get('note'):
                st.caption(schedule.get('note'))
            if schedule.get('submitted_at'):
                st.caption(f"Submitted: {schedule.get('submitted_at')}")
            review_note = st.text_area('Admin review note', value=schedule.get('review_note', ''), key=f"schedule_review_note_{item.get('saved_resume_id', '')}", height=100)
            c1, c2 = st.columns(2)
            with c1:
                if st.button('Mark useful', key=f"schedule_useful_{item.get('saved_resume_id', '')}", use_container_width=True):
                    _review_interview_schedule(item, 'useful', admin_user, review_note)
                    st.success('Interview schedule marked as useful.')
                    st.rerun()
            with c2:
                if st.button('Decline', key=f"schedule_decline_{item.get('saved_resume_id', '')}", use_container_width=True):
                    _review_interview_schedule(item, 'declined', admin_user, review_note)
                    st.success('Interview schedule declined.')
                    st.rerun()


def my_weekly_result_page(user: dict) -> None:
    show_header(user)
    st.subheader('My Weekly Result')
    rows = [row for row in _application_metrics_rows() if row.get('user_id') == user.get('id')]
    title_col, week_col = st.columns([3, 2])
    with title_col:
        st.markdown('**Weekly applications and interview schedules**')
    with week_col:
        selected_week_start = _metrics_week_selector(rows, key='my_weekly_result_week')

    summary_rows = _build_weekly_summary_rows(rows, [user], selected_week_start, include_openai=False)
    if summary_rows:
        st.dataframe(
            summary_rows,
            use_container_width=True,
            hide_index=True,
            column_config=_application_metrics_column_config(include_openai=False),
        )
    else:
        st.info('No saved applications for the selected week yet.')

    week_end = selected_week_start + timedelta(days=6)
    schedule_rows = [
        row for row in rows
        if row.get('schedule_submitted_at') and selected_week_start <= row['schedule_submitted_at'].date() <= week_end
    ]
    st.markdown('**Interview schedule submissions for the selected week**')
    if not schedule_rows:
        st.info('No interview schedules submitted in the selected week.')
    else:
        for row in sorted(schedule_rows, key=lambda item: item.get('schedule_submitted_at') or datetime.min, reverse=True):
            st.write(
                f"- {row.get('job_company', '') or 'Company'} — {row.get('job_title', '') or 'Role'} "
                f"({row.get('schedule_review_status', 'not_submitted')})"
            )

def user_access_page(user: dict) -> None:
    if not is_admin(user):
        st.error('Only admins can manage users.')
        return
    show_header(user)
    st.subheader('User Access')
    profiles = storage.get_profiles()
    users = storage.get_users()
    pending_users = [item for item in users if item.get('status') == 'pending']
    approved_users = [item for item in users if item.get('status') == 'approved']
    pending_tab, approved_tab, metrics_tab, schedule_reviews_tab = st.tabs(['Pending requests', 'Approved users', 'Application metrics', 'Interview schedule reviews'])
    with pending_tab:
        if not pending_users:
            st.info('No pending access requests.')
        for pending in pending_users:
            with st.expander(f"{pending.get('full_name') or pending.get('username')} • {pending.get('username')}"):
                available_profiles, owner_map = _available_profiles_for_user_assignment(profiles, users, pending.get('id', ''), pending.get('assigned_profile_ids', []))
                assigned = st.multiselect(
                    'Assign profiles on approval',
                    available_profiles,
                    default=[p for p in available_profiles if p.get('id') in (pending.get('assigned_profile_ids', []) or [])],
                    format_func=_format_profile_option,
                    key=f"pending_profiles_{pending.get('id')}"
                )
                st.caption(_assigned_profile_help_text(assigned, owner_map))
                make_admin = st.checkbox('Grant admin access', value=bool(pending.get('is_admin', False)), key=f"pending_admin_{pending.get('id')}")
                c1, c2 = st.columns(2)
                with c1:
                    if st.button('Approve', key=f"approve_user_{pending.get('id')}", use_container_width=True):
                        assigned_ids = [item.get('id') for item in assigned]
                        conflicting = [owner_map.get(str(profile_id)) for profile_id in assigned_ids if str(profile_id) in owner_map]
                        if conflicting:
                            st.error('One or more selected profiles are already assigned to another user. Refresh and choose only unassigned profiles.')
                        else:
                            storage.update_user(pending.get('id', ''), {'status': 'approved', 'is_admin': make_admin, 'assigned_profile_ids': assigned_ids, 'approved_at': datetime.utcnow().isoformat() + 'Z', 'approved_by_user_id': user.get('id', '')})
                            st.success('User approved.')
                            st.rerun()
                with c2:
                    if st.button('Reject / delete', key=f"reject_user_{pending.get('id')}", use_container_width=True):
                        remaining = [item for item in storage.get_users() if item.get('id') != pending.get('id')]
                        storage._write_json(storage.users_path, remaining)
                        st.success('Pending request removed.')
                        st.rerun()
    with approved_tab:
        for member in approved_users:
            with st.expander(f"{member.get('full_name') or member.get('username')} • {'Admin' if member.get('is_admin') else 'User'}"):
                available_profiles, owner_map = _available_profiles_for_user_assignment(profiles, users, member.get('id', ''), member.get('assigned_profile_ids', []))
                assigned = st.multiselect(
                    'Assigned profiles',
                    available_profiles,
                    default=[p for p in available_profiles if p.get('id') in (member.get('assigned_profile_ids', []) or [])],
                    format_func=_format_profile_option,
                    key=f"approved_profiles_{member.get('id')}"
                )
                st.caption(_assigned_profile_help_text(assigned, owner_map))
                member_is_admin = st.checkbox('Admin user', value=bool(member.get('is_admin', False)), key=f"approved_admin_{member.get('id')}")
                status = st.selectbox('Status', ['approved', 'disabled'], index=['approved', 'disabled'].index(member.get('status', 'approved') if member.get('status', 'approved') in ['approved', 'disabled'] else 'approved'), key=f"approved_status_{member.get('id')}")
                new_password = st.text_input('Reset password (optional)', type='password', key=f"reset_password_{member.get('id')}")
                if st.button('Save user access', key=f"save_user_{member.get('id')}", type='primary', use_container_width=True):
                    assigned_ids = [item.get('id') for item in assigned]
                    conflicting = [owner_map.get(str(profile_id)) for profile_id in assigned_ids if str(profile_id) in owner_map]
                    if conflicting:
                        st.error('One or more selected profiles are already assigned to another user. Refresh and choose only unassigned profiles.')
                    else:
                        patch = {'assigned_profile_ids': assigned_ids, 'is_admin': member_is_admin, 'status': status}
                        if new_password.strip():
                            patch |= build_password_record(new_password.strip()) | {'force_password_change': False}
                        storage.update_user(member.get('id', ''), patch)
                        st.success('User updated.')
                        st.rerun()

    with metrics_tab:
        _render_application_metrics_tab()

    with schedule_reviews_tab:
        _render_schedule_reviews_tab(user)

# ---------- Job list ----------

def _job_summary_label(job: dict) -> str:
    company = job.get('company', 'Unknown company') or 'Unknown company'
    title = job.get('job_title', 'Untitled role') or 'Untitled role'
    return f"{company} — {title} [{_region_label(job.get('region', ''))}]"


def _load_job_into_dashboard(job: dict) -> None:
    st.session_state['last_job_id'] = job.get('id', '')
    st.session_state['last_job_company'] = job.get('company', '')
    st.session_state['last_job_link'] = job.get('link', '')
    st.session_state['last_target_role'] = job.get('job_title', '')
    st.session_state['last_job_description'] = job.get('description', '')
    st.session_state['last_job_region'] = _normalize_region(job.get('region', ''))
    queue_nav('Dashboard')


def _extract_url_from_line(line: str) -> str:
    match = re.search(r'https?://\S+', line)
    return match.group(0).strip() if match else ''


def _parse_batch_jobs(text: str, user: dict) -> list[dict]:
    jobs: list[dict] = []
    for raw_line in [line.strip() for line in str(text or '').splitlines() if line.strip()]:
        parts = [part.strip() for part in raw_line.split('|')]
        link = ''
        company = ''
        job_title = ''
        note = ''
        region = 'US'
        if len(parts) >= 3 and parts[2].startswith('http'):
            company = parts[0]
            job_title = parts[1]
            link = parts[2]
            if len(parts) > 3:
                maybe_region = _normalize_region(parts[3])
                if maybe_region in {'ANY', 'US', 'EU', 'LATAM'}:
                    region = maybe_region
                    note = parts[4] if len(parts) > 4 else ''
                else:
                    note = parts[3]
        else:
            link = _extract_url_from_line(raw_line)
            note = raw_line.replace(link, '').strip(' -|') if link else raw_line
            if len(parts) >= 2 and not link:
                company = parts[0]
                job_title = parts[1]
                if len(parts) > 2:
                    maybe_region = _normalize_region(parts[2])
                    if maybe_region in {'ANY', 'US', 'EU', 'LATAM'}:
                        region = maybe_region
                        note = parts[3] if len(parts) > 3 else ''
                    else:
                        note = parts[2]
        jobs.append({
            'id': storage.make_id('job'),
            'company': company,
            'job_title': job_title,
            'description': '',
            'link': link,
            'region': region,
            'note': note,
            'status': 'pending',
            'source': 'batch',
            'scrape_status': 'queued' if link else 'done',
            'scrape_error': '',
            'created_by_user_id': user.get('id', ''),
            'created_by_username': user.get('username', ''),
            'submitted_at': datetime.utcnow().isoformat() + 'Z',
            'approved_at': '',
            'approved_by_user_id': '',
            'approved_by_username': '',
        })
    return jobs


def _job_scrape_progress_scope(user: dict) -> list[dict]:
    jobs = storage.get_jobs(include_pending=True)
    if is_admin(user):
        return [job for job in jobs if job.get('status') == 'pending']
    return [
        job for job in jobs
        if job.get('status') == 'pending' and job.get('created_by_user_id') == user.get('id', '')
    ]


def _render_job_scrape_progress(user: dict, title: str = 'Background scrape progress') -> None:
    scoped_jobs = _job_scrape_progress_scope(user)
    queued = sum(1 for job in scoped_jobs if job.get('scrape_status') == 'queued')
    processing = sum(1 for job in scoped_jobs if job.get('scrape_status') == 'processing')
    done = sum(1 for job in scoped_jobs if job.get('scrape_status') == 'done')
    errors = sum(1 for job in scoped_jobs if job.get('scrape_status') == 'error')
    total = len(scoped_jobs)
    if total == 0:
        st.caption('No pending background scraping jobs.')
        return
    st.caption(f"{title}: {done} done • {processing} processing • {queued} queued • {errors} errors")
    st.progress((done + errors) / total if total else 0.0, text=f'{done + errors} of {total} completed')


def job_list_page(user: dict) -> None:
    show_header(user)
    st.subheader('Job List')
    accessible_profiles = get_accessible_profiles(user)
    generated_items = storage.get_generated_resumes()
    if not is_admin(user):
        generated_items = [item for item in generated_items if item.get('created_by_user_id') == user.get('id')]
    applied_map: dict[str, set[str]] = {}
    for item in generated_items:
        job_id = str(item.get('job_id', '')).strip()
        profile_id = str(item.get('profile_id', '')).strip()
        if job_id and profile_id:
            applied_map.setdefault(job_id, set()).add(profile_id)

    notice = st.session_state.pop('job_list_notice', '')
    if notice:
        st.success(notice)

    tabs = ['Approved jobs', 'Add job'] + (['Batch presave', 'Pending admin queue', 'Reported jobs'] if is_admin(user) else [])
    rendered_tabs = st.tabs(tabs)

    with rendered_tabs[0]:
        search_text = st.text_input('Search approved jobs', placeholder='Search by company, role, or note')
        approved_jobs = storage.get_jobs(include_pending=False)
        if accessible_profiles:
            approved_jobs = [job for job in approved_jobs if any(_profile_matches_job_region(profile, job) for profile in accessible_profiles)]
        filtered_jobs = []
        needle = search_text.strip().lower()
        for job in approved_jobs:
            blob = ' '.join([job.get('company', ''), job.get('job_title', ''), job.get('description', ''), job.get('note', ''), job.get('region', '')]).lower()
            if not needle or needle in blob:
                filtered_jobs.append(job)
        if not filtered_jobs:
            st.info('No approved jobs match your assigned profile markets yet.')
        for job in filtered_jobs:
            with st.container(border=True):
                info_col, action_col = st.columns([5.4, 1.2], gap='medium')
                with info_col:
                    st.markdown(f"**{_job_summary_label(job)}**")
                    if job.get('link'):
                        st.caption(job.get('link', ''))
                    matching_profiles = [profile for profile in accessible_profiles if _profile_matches_job_region(profile, job)]
                    if matching_profiles:
                        applied_bits = []
                        for profile in matching_profiles:
                            icon = '✅' if profile.get('id') in applied_map.get(job.get('id', ''), set()) else '⬜'
                            applied_bits.append(f"{icon} {_format_profile_option(profile)}")
                        if applied_bits:
                            st.caption('   '.join(applied_bits))
                    else:
                        st.caption('No assigned profiles match this job market.')
                    meta_bits = [f"Market: {_region_label(job.get('region', ''))}"]
                    if job.get('created_by_username'):
                        meta_bits.append(f"Added by {job.get('created_by_username', '')}")
                    if job.get('note'):
                        meta_bits.append(job.get('note', ''))
                    if meta_bits:
                        st.caption(' • '.join(meta_bits))
                with action_col:
                    if st.button('Use in Dashboard', key=f"use_job_{job.get('id')}", use_container_width=True):
                        _load_job_into_dashboard(job)
                        st.rerun()
                    if is_admin(user):
                        edit_key = f'edit_job_open_{job.get("id")}'
                        confirm_key = f'delete_job_confirm_{job.get("id")}'
                        if st.button('Edit', key=f"edit_job_{job.get('id')}", use_container_width=True):
                            st.session_state[edit_key] = not st.session_state.get(edit_key, False)
                        if st.session_state.get(confirm_key, False):
                            d_yes, d_no = st.columns(2)
                            with d_yes:
                                if st.button('Confirm', key=f"delete_job_yes_{job.get('id')}", type='primary', use_container_width=True):
                                    storage.delete_job(job.get('id', ''))
                                    st.session_state.pop(confirm_key, None)
                                    st.session_state.pop(edit_key, None)
                                    st.session_state['job_list_notice'] = 'Job deleted.'
                                    st.rerun()
                            with d_no:
                                if st.button('Cancel', key=f"delete_job_no_{job.get('id')}", use_container_width=True):
                                    st.session_state.pop(confirm_key, None)
                                    st.rerun()
                        else:
                            if st.button('Delete', key=f"delete_job_{job.get('id')}", use_container_width=True):
                                st.session_state[confirm_key] = True
                                st.rerun()
                if is_admin(user) and st.session_state.get(f'edit_job_open_{job.get("id")}', False):
                    with st.form(key=f"approved_edit_form_{job.get('id')}"):
                        ec_company = st.text_input('Company', value=job.get('company', ''), key=f"approved_edit_company_{job.get('id')}")
                        ec_title = st.text_input('Job title', value=job.get('job_title', ''), key=f"approved_edit_title_{job.get('id')}")
                        ec_link = st.text_input('Link', value=job.get('link', ''), key=f"approved_edit_link_{job.get('id')}")
                        region_label = _region_label(job.get('region', 'US'))
                        try:
                            region_index = REGION_OPTIONS.index(region_label)
                        except ValueError:
                            region_index = REGION_OPTIONS.index('US') if 'US' in REGION_OPTIONS else 0
                        ec_region = st.selectbox('Job market', REGION_OPTIONS, index=region_index, key=f"approved_edit_region_{job.get('id')}")
                        ec_desc = st.text_area('Description', value=job.get('description', ''), height=240, key=f"approved_edit_desc_{job.get('id')}")
                        ec_note = st.text_area('Note', value=job.get('note', ''), height=90, key=f"approved_edit_note_{job.get('id')}")
                        ec_save, ec_cancel = st.columns(2)
                        save_clicked = ec_save.form_submit_button('Save changes', type='primary', use_container_width=True)
                        cancel_clicked = ec_cancel.form_submit_button('Close', use_container_width=True)
                    if save_clicked:
                        storage.update_job(job.get('id', ''), {
                            'company': ec_company,
                            'job_title': ec_title,
                            'link': ec_link,
                            'region': _normalize_region(ec_region),
                            'description': ec_desc,
                            'note': ec_note,
                        })
                        st.session_state[f'edit_job_open_{job.get("id")}'] = False
                        st.session_state['job_list_notice'] = 'Job updated.'
                        st.rerun()
                    if cancel_clicked:
                        st.session_state[f'edit_job_open_{job.get("id")}'] = False
                        st.rerun()

    with rendered_tabs[1]:
        with st.form('add_job_form'):
            company = st.text_input('Company')
            job_title = st.text_input('Job title / role')
            c_add1, c_add2 = st.columns(2)
            with c_add1:
                link = st.text_input('Link')
            with c_add2:
                region = st.selectbox('Job market', REGION_OPTIONS, index=REGION_OPTIONS.index('US'))
            description = st.text_area('Job description', height=220)
            note = st.text_area('Note', height=100)
            confirm_duplicate = st.checkbox('I confirm I still want to add this if the same company and role already exist')
            submitted = st.form_submit_button('Add job', type='primary')
        if submitted:
            duplicate = storage.find_duplicate_job(company, job_title)
            if duplicate and not confirm_duplicate:
                st.warning(f"A matching job already exists: {_job_summary_label(duplicate)}. Tick the confirmation box to add it anyway.")
                return
            status = 'approved' if is_admin(user) else 'pending'
            scrape_status = 'done' if description.strip() else ('queued' if link.strip() else 'done')
            job_payload = {
                'id': storage.make_id('job'),
                'company': company.strip(),
                'job_title': job_title.strip(),
                'description': description.strip(),
                'link': link.strip(),
                'region': _normalize_region(region),
                'note': note.strip(),
                'status': status,
                'source': 'manual',
                'scrape_status': scrape_status,
                'scrape_error': '',
                'created_by_user_id': user.get('id', ''),
                'created_by_username': user.get('username', ''),
                'submitted_at': datetime.utcnow().isoformat() + 'Z',
                'approved_at': datetime.utcnow().isoformat() + 'Z' if status == 'approved' else '',
                'approved_by_user_id': user.get('id', '') if status == 'approved' else '',
                'approved_by_username': user.get('username', '') if status == 'approved' else '',
            }
            storage.upsert_job(job_payload)
            st.session_state['job_list_notice'] = 'Job saved.' if status == 'approved' else 'Job submitted for admin approval.'
            st.rerun()

    if is_admin(user):
        with rendered_tabs[2]:
            st.caption('Paste one job per line. Supported formats: URL only, Company | Role | URL | Note, or Company | Role | URL | Region | Note. Default market is US when not provided.')
            _render_job_scrape_progress(user)
            batch_text = st.text_area('Batch job input', height=220, key='batch_jobs_text')
            if st.button('Queue batch jobs', type='primary', use_container_width=True):
                jobs = _parse_batch_jobs(batch_text, user)
                if not jobs:
                    st.error('Paste at least one job line first.')
                else:
                    storage.bulk_upsert_jobs(jobs)
                    st.session_state['job_list_notice'] = f'{len(jobs)} job entries queued for background scraping and admin review.'
                    st.rerun()

        with rendered_tabs[3]:
            pending_jobs = [job for job in storage.get_jobs(include_pending=True) if job.get('status') == 'pending']
            header_col, action_col = st.columns([4.5, 1.5])
            with header_col:
                _render_job_scrape_progress(user, title='Pending queue progress')
            with action_col:
                st.write('')
                st.write('')
                if pending_jobs and st.button('Approve all job lists', type='primary', use_container_width=True):
                    approve_patch = {}
                    approved_at = datetime.utcnow().isoformat() + 'Z'
                    for job in pending_jobs:
                        approve_patch[job.get('id', '')] = {
                            'status': 'approved',
                            'approved_at': approved_at,
                            'approved_by_user_id': user.get('id', ''),
                            'approved_by_username': user.get('username', ''),
                            'scrape_status': 'done' if str(job.get('description', '')).strip() else job.get('scrape_status', 'queued'),
                        }
                    storage.bulk_update_jobs(approve_patch)
                    st.session_state['job_list_notice'] = f'{len(approve_patch)} pending jobs approved.'
                    st.rerun()
            if not pending_jobs:
                st.info('No pending jobs to review.')
            else:
                st.caption('Pending job edits are grouped in submit forms so typing does not rerender the whole page.')
            for job in pending_jobs:
                with st.expander(f"Pending • {_job_summary_label(job)}"):
                    st.caption(f"Source: {job.get('source', 'manual')} • Scrape status: {job.get('scrape_status', 'n/a')}")
                    if job.get('scrape_error'):
                        st.warning(job.get('scrape_error'))
                    with st.form(key=f"pending_job_form_{job.get('id')}", clear_on_submit=False):
                        company_value = st.text_input('Company', value=job.get('company', ''), key=f"pending_job_company_form_{job.get('id')}")
                        title_value = st.text_input('Job title', value=job.get('job_title', ''), key=f"pending_job_title_form_{job.get('id')}")
                        link_value = st.text_input('Link', value=job.get('link', ''), key=f"pending_job_link_form_{job.get('id')}")
                        region_label = _region_label(job.get('region', 'US'))
                        try:
                            region_index = REGION_OPTIONS.index(region_label)
                        except ValueError:
                            region_index = REGION_OPTIONS.index('US') if 'US' in REGION_OPTIONS else 0
                        region_value = st.selectbox('Job market', REGION_OPTIONS, index=region_index, key=f"pending_job_region_form_{job.get('id')}")
                        desc_value = st.text_area('Description', value=job.get('description', ''), height=220, key=f"pending_job_desc_form_{job.get('id')}")
                        note_value = st.text_area('Note', value=job.get('note', ''), height=90, key=f"pending_job_note_form_{job.get('id')}")
                        c1, c2, c3, c4 = st.columns(4)
                        save_clicked = c1.form_submit_button('Save draft', use_container_width=True)
                        approve_clicked = c2.form_submit_button('Approve', type='primary', use_container_width=True)
                        requeue_clicked = c3.form_submit_button('Requeue scrape', use_container_width=True)
                        delete_clicked = c4.form_submit_button('Delete', use_container_width=True)

                    base_patch = {
                        'company': company_value,
                        'job_title': title_value,
                        'link': link_value,
                        'region': _normalize_region(region_value),
                        'description': desc_value,
                        'note': note_value,
                    }
                    if save_clicked:
                        storage.update_job(job.get('id', ''), base_patch)
                        st.session_state['job_list_notice'] = 'Pending job updated.'
                        st.rerun()
                    if approve_clicked:
                        patch = dict(base_patch)
                        patch.update({
                            'status': 'approved',
                            'approved_at': datetime.utcnow().isoformat() + 'Z',
                            'approved_by_user_id': user.get('id', ''),
                            'approved_by_username': user.get('username', ''),
                            'scrape_status': 'done' if str(desc_value).strip() else job.get('scrape_status', 'queued'),
                        })
                        storage.update_job(job.get('id', ''), patch)
                        st.session_state['job_list_notice'] = 'Job approved and visible to all users.'
                        st.rerun()
                    if requeue_clicked:
                        patch = dict(base_patch)
                        patch.update({'scrape_status': 'queued', 'scrape_error': ''})
                        storage.update_job(job.get('id', ''), patch)
                        st.session_state['job_list_notice'] = 'Job queued for background scraping.'
                        st.rerun()
                    if delete_clicked:
                        storage.delete_job(job.get('id', ''))
                        st.session_state['job_list_notice'] = 'Pending job deleted.'
                        st.rerun()

        with rendered_tabs[4]:
            reported_jobs = [
                job for job in storage.get_jobs(include_pending=True)
                if ((job.get('reports') or []) or job.get('flagged'))
                and not job.get('admin_applied', False)
            ]
            reported_jobs.sort(
                key=lambda job: ((job.get('reports') or [{}])[-1].get('reported_at', '')),
                reverse=True,
            )
            if not reported_jobs:
                st.info('No reported jobs yet.')
            else:
                st.caption(f'{len(reported_jobs)} reported job(s). Reports are flagged by users or auto-flagged after repeated low ATS scores.')
            for job in reported_jobs:
                reports = job.get('reports', []) or []
                with st.expander(f"Reported • {_job_summary_label(job)} • {len(reports)} report(s)"):
                    if job.get('link'):
                        st.caption(job.get('link', ''))
                    meta_bits = [f"Status: {job.get('status', 'unknown')}", f"Market: {_region_label(job.get('region', ''))}"]
                    if job.get('created_by_username'):
                        meta_bits.append(f"Added by {job.get('created_by_username', '')}")
                    st.caption(' • '.join(meta_bits))
                    if not reports:
                        st.warning('Job is flagged but has no report entries.')
                    for index, report in enumerate(reports):
                        source = str(report.get('source', 'user') or 'user').strip()
                        reporter = report.get('reported_by_username') or ('system' if source == 'system' else 'unknown')
                        reported_at = report.get('reported_at', '') or 'unknown time'
                        st.markdown(f"**{index + 1}. {reporter}** ({source}) — `{reported_at}`")
                        st.write(report.get('reason', ''))
                    action1, action2, action3 = st.columns(3)
                    with action1:
                        if st.button('Dismiss reports & restore', key=f'dismiss_reports_{job.get("id")}', use_container_width=True):
                            storage.clear_job_reports(job.get('id', ''))
                            st.session_state['job_list_notice'] = 'Reports cleared and job restored to the active list.'
                            st.rerun()
                    with action2:
                        if st.button('Delete job', key=f'delete_reported_{job.get("id")}', use_container_width=True):
                            storage.delete_job(job.get('id', ''))
                            st.session_state['job_list_notice'] = 'Reported job deleted.'
                            st.rerun()
                    with action3:
                        if st.button('Mark as applied by admin', key=f'admin_applied_{job.get("id")}', use_container_width=True):
                            storage.update_job(job.get('id', ''), {
                                'admin_applied': True,
                                'admin_applied_at': datetime.utcnow().isoformat() + 'Z',
                                'admin_applied_by_user_id': user.get('id', ''),
                                'admin_applied_by_username': user.get('username', ''),
                                'flagged': False,
                            })
                            st.session_state['job_list_notice'] = 'Job marked as applied by admin and removed from active queue.'
                            st.rerun()


# ---------- Generated resumes ----------

def _resume_application_date_label(created_at: str) -> str:
    value = str(created_at or '').strip()
    if not value:
        return 'Unknown date'
    try:
        return datetime.fromisoformat(value.replace('Z', '+00:00')).strftime('%Y-%m-%d')
    except Exception:
        return value.split('T')[0] if 'T' in value else value[:10]


def _generated_resume_display_title(item: dict) -> str:
    date_label = _resume_application_date_label(item.get('created_at', ''))
    title = item.get('job_title') or item.get('target_role') or item.get('resume', {}).get('headline') or 'Untitled role'
    return f'{date_label} • {title}'


def _generated_resume_search_blob(item: dict, profile_name: str) -> str:
    parts = [
        item.get('job_company', ''),
        item.get('job_title', ''),
        item.get('target_role', ''),
        profile_name,
        item.get('created_by_username', ''),
        item.get('company_message', ''),
        item.get('job_link', ''),
        item.get('job_description', ''),
        item.get('job_region', ''),
        json.dumps(item.get('resume', {}) or {}, ensure_ascii=False),
    ]
    return ' '.join(str(part) for part in parts if part).lower()


def _render_generated_resume_download_tab(item: dict, resume_snapshot: dict, item_key: str) -> None:
    profile = storage.get_profile_by_id(item.get('profile_id', '')) or {
        'name': item.get('resume', {}).get('name', '') or item.get('created_by_username', 'Candidate'),
        'email': '',
        'phone': '',
        'location': '',
        'linkedin': '',
        'github': '',
        'uploaded_resume': {},
    }
    filename = item.get('download_filename') or item.get('saved_files', {}).get('pdf') or f"{_build_file_stem(profile)}.pdf"
    st.write(f"Download filename: {filename}")
    try:
        exports = _build_uploaded_docx_pdf_exports(resume=resume_snapshot, profile=profile, app_settings=storage.get_app_settings())
    except Exception as exc:
        st.error(str(exc))
        return
    st.download_button(
        'Download Resume PDF',
        data=exports.get('pdf', b''),
        file_name=filename,
        mime='application/pdf',
        use_container_width=True,
        disabled=not bool(exports.get('pdf')),
        key=f'generated_resume_download_{item_key}',
    )
    st.caption('The PDF is regenerated from the saved resume snapshot and the uploaded profile DOCX style.')

def _render_interview_schedule_tab(item: dict, item_key: str, user: dict) -> None:
    schedule = item.get('interview_schedule', {}) if isinstance(item.get('interview_schedule', {}), dict) else {}
    status = str(schedule.get('review_status', 'not_submitted') or 'not_submitted').strip()
    status_map = {
        'not_submitted': ('Not submitted', st.info),
        'waiting_review': ('Waiting for review', st.warning),
        'useful': ('Approved as useful', st.success),
        'declined': ('Declined', st.error),
    }
    status_label, status_writer = status_map.get(status, (status.replace('_', ' ').title(), st.info))
    status_writer(f'Status: {status_label}')
    if schedule.get('reviewed_by_username'):
        st.caption(f"Reviewed by {schedule.get('reviewed_by_username')} on {schedule.get('reviewed_at', '') or 'n/a'}")
    if schedule.get('review_note'):
        st.caption(f"Review note: {schedule.get('review_note')}")

    interviewer_key = f"schedule_interviewer_{item_key}"
    interview_time_key = f"schedule_time_{item_key}"
    meeting_link_key = f"schedule_link_{item_key}"
    note_key = f"schedule_note_{item_key}"
    if interviewer_key not in st.session_state:
        st.session_state[interviewer_key] = schedule.get('interviewer_name', '')
    if interview_time_key not in st.session_state:
        st.session_state[interview_time_key] = schedule.get('interview_time', '')
    if meeting_link_key not in st.session_state:
        st.session_state[meeting_link_key] = schedule.get('meeting_link', '')
    if note_key not in st.session_state:
        st.session_state[note_key] = schedule.get('note', '')

    st.text_input('Interviewer name', key=interviewer_key, placeholder='e.g. John Smith')
    st.text_input('Interview time', key=interview_time_key, placeholder='e.g. 2026-04-25 14:00 PST')
    st.text_input('Meeting link', key=meeting_link_key, placeholder='https://...')
    st.text_area('Schedule note (optional)', key=note_key, height=100)

    if schedule.get('submitted_at'):
        st.caption(f"Last submitted: {schedule.get('submitted_at')}")

    submit_label = 'Submit interview schedule for review' if status in {'not_submitted', ''} else 'Update and resubmit for review'
    if st.button(submit_label, key=f"submit_schedule_{item_key}", use_container_width=True):
        interviewer_name = st.session_state.get(interviewer_key, '').strip()
        interview_time = st.session_state.get(interview_time_key, '').strip()
        meeting_link = st.session_state.get(meeting_link_key, '').strip()
        note = st.session_state.get(note_key, '').strip()
        if not interviewer_name or not interview_time or not meeting_link:
            st.error('Interviewer name, interview time, and meeting link are required.')
        else:
            _submit_interview_schedule(item, interviewer_name, interview_time, meeting_link, note)
            st.success('Interview schedule submitted and marked as waiting for review.')
            st.rerun()


def generated_resumes_page(user: dict) -> None:
    show_header(user)
    items = storage.get_generated_resumes()
    admin_view = is_admin(user)
    if not admin_view:
        items = [item for item in items if item.get('created_by_user_id') == user.get('id')]
    if not items:
        st.subheader('Generated resumes')
        st.info('No generated resumes yet.')
        return

    profiles_map = {item.get('id'): item for item in storage.get_profiles()}
    company_options = sorted({str(item.get('job_company', '')).strip() for item in items if str(item.get('job_company', '')).strip()})
    profile_options = sorted({str((profiles_map.get(item.get('profile_id')) or {}).get('name', '')).strip() for item in items if str((profiles_map.get(item.get('profile_id')) or {}).get('name', '')).strip()})

    default_date = st.session_state.get('generated_resume_filter_date')
    if not isinstance(default_date, date):
        default_date = date.today()
    default_start = st.session_state.get('generated_resume_filter_start_date')
    if not isinstance(default_start, date):
        default_start = default_date
    default_end = st.session_state.get('generated_resume_filter_end_date')
    if not isinstance(default_end, date):
        default_end = default_date

    title_placeholder = st.empty()
    if admin_view:
        f1, f2, f3, f4, f5 = st.columns([2.0, 1.0, 1.0, 0.9, 0.9])
    else:
        f1, f2, f3, f4 = st.columns([2.1, 1.0, 1.0, 1.0])
    with f1:
        search_text = st.text_input('Search applied resumes', placeholder='Search by company, job title, profile, application message, or job description', key='generated_resume_search_text')
    with f2:
        selected_company = st.selectbox('Company', ['All'] + company_options, key='generated_resume_search_company')
    with f3:
        selected_profile = st.selectbox('Profile', ['All'] + profile_options, key='generated_resume_search_profile')
    if admin_view:
        with f4:
            selected_start_date = st.date_input('From', value=default_start, key='generated_resume_filter_start_date')
        with f5:
            selected_end_date = st.date_input('To', value=default_end, key='generated_resume_filter_end_date')
        if selected_start_date > selected_end_date:
            selected_start_date, selected_end_date = selected_end_date, selected_start_date
            st.session_state['generated_resume_filter_start_date'] = selected_start_date
            st.session_state['generated_resume_filter_end_date'] = selected_end_date
    else:
        with f4:
            selected_date = st.date_input('Application date', value=default_date, key='generated_resume_filter_date')

    filtered_items = []
    needle = str(search_text or '').strip().lower()
    selected_date_str = ''
    selected_start_str = ''
    selected_end_str = ''
    if admin_view:
        selected_start_str = selected_start_date.strftime('%Y-%m-%d') if isinstance(selected_start_date, date) else ''
        selected_end_str = selected_end_date.strftime('%Y-%m-%d') if isinstance(selected_end_date, date) else ''
    else:
        selected_date_str = selected_date.strftime('%Y-%m-%d') if isinstance(selected_date, date) else ''

    for item in items:
        profile_name = str((profiles_map.get(item.get('profile_id')) or {}).get('name', '')).strip()
        if selected_company != 'All' and str(item.get('job_company', '')).strip() != selected_company:
            continue
        if selected_profile != 'All' and profile_name != selected_profile:
            continue
        created_label = _resume_application_date_label(item.get('created_at', ''))
        if admin_view:
            if selected_start_str and created_label < selected_start_str:
                continue
            if selected_end_str and created_label > selected_end_str:
                continue
        else:
            if selected_date_str and created_label != selected_date_str:
                continue
        if needle and needle not in _generated_resume_search_blob(item, profile_name):
            continue
        filtered_items.append(item)

    filtered_count = len(filtered_items)
    if admin_view:
        range_label = f'{selected_start_str or "the start"} → {selected_end_str or "the end"}'
        title_placeholder.subheader(f'Generated resumes • {filtered_count} in {range_label}')
        st.caption(f'Showing {filtered_count} of {len(items)} saved resumes for {range_label}.')
    else:
        title_placeholder.subheader(f'Generated resumes • {filtered_count} on {selected_date_str or "the selected date"}')
        st.caption(f'Showing {filtered_count} of {len(items)} saved resumes for {selected_date_str or "the selected date"}.')
    if not filtered_items:
        st.info('No saved resume matched your filters.')
        return

    use_ai = st.toggle('Use OpenAI for application answers', value=True, key='generated_resume_use_ai')
    for index, item in enumerate(reversed(filtered_items), start=1):
        created_at = item.get('created_at', '')
        item_key = f"{index}_{created_at}_{item.get('saved_resume_id', '')}"
        resume_snapshot = item.get('resume', {}) or {}
        job_description = item.get('job_description', '') or ''
        target_role = item.get('target_role', '') or ''
        profile_name = str((profiles_map.get(item.get('profile_id')) or {}).get('name', '')).strip() or 'Unknown profile'
        title = _generated_resume_display_title(item)
        with st.expander(title):
            meta_cols = st.columns([1.25, 1.25, 1, 1])
            with meta_cols[0]:
                st.caption('Company')
                st.write(item.get('job_company', '') or '—')
            with meta_cols[1]:
                st.caption('Profile')
                st.write(profile_name)
            with meta_cols[2]:
                st.caption('ATS at save')
                st.write(f"{item.get('ats_score', 'n/a')}/100" if item.get('ats_score') not in [None, ''] else 'n/a')
            with meta_cols[3]:
                st.caption('Created by')
                st.write(item.get('created_by_username', 'n/a'))

            saved_tab_labels = ['Snapshot', 'ATS Score']
            if is_admin(user):
                saved_tab_labels.append('Job Application Answers')
            saved_tab_labels.extend(['Job Description', 'Interview Schedule', 'Download Resume'])
            saved_tabs = dict(zip(saved_tab_labels, st.tabs(saved_tab_labels)))
            with saved_tabs['Snapshot']:
                st.write(f"Created: {created_at or 'n/a'}")
                st.write(f"Resume ID: {item.get('saved_resume_id', 'n/a')}")
                st.write("Style source: Uploaded DOCX")
                st.write(f"Company message status: {item.get('company_message_status', 'n/a')}")
                if item.get('job_link'):
                    st.write(f"Job link: {item.get('job_link')}")
                company_message_key = f"company_message_snapshot_{item_key}"
                if company_message_key not in st.session_state:
                    st.session_state[company_message_key] = item.get('company_message', '')
                st.text_area('Company message / application email', key=company_message_key, height=160)
                if st.button('Save company message', key=f"save_company_message_{item_key}", use_container_width=True):
                    company_message_value = st.session_state.get(company_message_key, '').strip()
                    if not company_message_value:
                        st.error('Paste the company message first.')
                    else:
                        _update_saved_resume_message(item, company_message_value)
                        st.success('Company message saved.')
                        st.rerun()
            with saved_tabs['ATS Score']:
                analysis = analyze_ats_score(resume_snapshot, job_description, target_role=target_role)
                _render_ats_analysis(analysis)
            if 'Job Application Answers' in saved_tabs:
                with saved_tabs['Job Application Answers']:
                    _render_application_answers_tab(
                        resume_snapshot=resume_snapshot,
                        job_description=job_description,
                        target_role=target_role,
                        use_ai=use_ai,
                        cache_prefix=f'generated_resume_answers_{item_key}',
                    )
            with saved_tabs['Job Description']:
                st.text_area('Job description', value=job_description, height=320, key=f"job_description_snapshot_{item_key}")
            with saved_tabs['Interview Schedule']:
                _render_interview_schedule_tab(item, item_key, user)
            with saved_tabs['Download Resume']:
                _render_generated_resume_download_tab(item, resume_snapshot, item_key)


# ---------- ATS / analysis ----------

def _render_ats_analysis(analysis: dict) -> None:
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric('ATS score', f"{analysis.get('overall_score', 0)}/100")
    with c2:
        st.metric('Matched JD keywords', str(len(analysis.get('matched_keywords', []))))
    with c3:
        st.metric('Missing JD keywords', str(len(analysis.get('missing_keywords', []))))
    st.caption('This is a local ATS-fit estimate based on keyword coverage, title alignment, skills alignment, experience evidence, and structure. It is designed for editing guidance, not as a claim about any specific ATS vendor.')
    s1, s2 = st.columns(2)
    with s1:
        st.markdown('**Strengths**')
        for line in analysis.get('strengths', []):
            st.write(f'- {line}')
        st.markdown('**Matched keywords**')
        st.write(', '.join(analysis.get('matched_keywords', [])) or 'None detected yet.')
    with s2:
        st.markdown('**Risks**')
        for line in analysis.get('risks', []):
            st.write(f'- {line}')
        st.markdown('**Missing keywords**')
        st.write(', '.join(analysis.get('missing_keywords', [])) or 'No major gaps detected.')
    st.markdown('**Category scores**')
    for category, score in analysis.get('category_scores', {}).items():
        st.write(f'- {category}: {score}')
    st.markdown('**Suggestions to improve this draft**')
    for line in analysis.get('suggestions', []):
        st.write(f'- {line}')


def _ats_notes_context_block(resume: dict) -> None:
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('**AI-tailored headline**')
        st.write(resume.get('headline', ''))
        st.markdown('**Fit keywords**')
        st.write(', '.join(resume.get('fit_keywords', [])) or 'No keywords extracted.')
    with c2:
        st.markdown('**Company role positioning**')
        for job in resume.get('work_history', []):
            st.write(f"- {job.get('company_name', '')}: {job.get('role_title', '')}")
            if job.get('role_headline'):
                st.caption(job.get('role_headline', ''))
    st.markdown('**Grouped skills**')
    for group in resume.get('skill_groups', []):
        items = ', '.join(group.get('items', []))
        if items:
            st.write(f"- {group.get('category', 'Other')}: {items}")


def _dashboard_ats_notes_tab(profile: dict, template: dict, resume: dict, job_description: str, target_role: str, custom_prompt: str, default_prompt: str, use_ai: bool, clean_generation: bool) -> None:
    if not str(job_description).strip():
        st.info('Add a job description to see ATS analysis.')
        return
    analysis = analyze_ats_score(resume, job_description, target_role=target_role)
    _render_ats_analysis(analysis)
    st.markdown('---')
    _ats_notes_context_block(resume)
    st.markdown('---')
    st.markdown('**Auto-improve resume to ATS target**')
    st.caption('This runs a short revision loop using the current visible draft, the ATS gaps, and your optional extra requirements. It stops once the score reaches the target or the round limit is hit.')
    st.text_area('Additional ATS improvement requirements (optional)', key='dashboard_ats_improve_prompt', height=110, placeholder='Example: Keep the summary concise, make bullets mention exact backend stacks, and avoid generic phrases.')
    a1, a2 = st.columns([1.1, 1])
    with a1:
        target_score = st.slider('Target ATS score', min_value=90, max_value=99, value=91, key='dashboard_ats_target_score')
    with a2:
        max_rounds = st.number_input('Max improvement rounds', min_value=1, max_value=5, value=3, step=1, key='dashboard_ats_max_rounds')
    if st.button('Auto-improve resume to ATS target', type='primary', use_container_width=True, key='dashboard_ats_improve_button'):
        current_resume = st.session_state.get('last_resume') or {}
        if not current_resume:
            st.error('Generate a resume first.')
            return
        with st.spinner('Improving the current draft against ATS guidance...'):
            result = improve_resume_to_target_ats(profile=profile, job_description=job_description, current_resume=current_resume, target_score=int(target_score), max_rounds=int(max_rounds), additional_requirements=st.session_state.get('dashboard_ats_improve_prompt', ''), target_role=target_role, custom_prompt=custom_prompt, default_prompt=default_prompt, use_ai=use_ai, clean_generation=clean_generation)
            _record_openai_usage_for_improve(result)
        updated_resume = result.get('resume', current_resume)
        exports = _build_uploaded_docx_pdf_exports(resume=updated_resume, profile=profile, app_settings=storage.get_app_settings())
        st.session_state['last_resume'] = updated_resume
        st.session_state['last_exports'] = exports
        st.session_state['last_generator_mode'] = result.get('mode', 'ats-improve')
        st.session_state['last_ats_improve_history'] = result.get('history', [])
        improved_ats = analyze_ats_score(updated_resume, job_description, target_role=target_role)
        improved_score = int((improved_ats or {}).get('overall_score', 0))
        improved_job_id = str(st.session_state.get('last_job_id', '') or '').strip()
        active_user = storage.get_user_by_id(st.session_state.get('current_user_id', '')) or {}
        if improved_job_id and _enforce_low_ats_rate_limit(active_user, improved_job_id, improved_score):
            st.rerun()
            return
        _queue_editor_reload(updated_resume, 'ATS-guided improvement applied to the current draft.')
        st.rerun()
    history = st.session_state.get('last_ats_improve_history', [])
    if history:
        st.markdown('**Recent ATS improvement rounds**')
        for item in history:
            st.write(f"- Round {item.get('round')}: {item.get('before_score')} → {item.get('after_score')} ({item.get('mode')})")
            for line in item.get('used_suggestions', [])[:3]:
                st.caption(line)


# ---------- Serialization helpers ----------

def _serialize_work_history(items: list[dict]) -> str:
    blocks: list[str] = []
    for item in items:
        header = ' | '.join([item.get('company_name', ''), item.get('duration', ''), item.get('location', '')]).strip()
        bullets = '\n'.join(f"- {bullet}" for bullet in item.get('bullets', []))
        blocks.append(f"{header}\n{bullets}".strip())
    return '\n\n'.join(blocks)


def _parse_work_history(value: str) -> list[dict]:
    items: list[dict] = []
    for block in [part.strip() for part in str(value or '').split('\n\n') if part.strip()]:
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        header = lines[0]
        parts = [part.strip() for part in header.split('|')]
        while len(parts) < 3:
            parts.append('')
        bullets = [line.lstrip('- ').strip() for line in lines[1:] if line.strip().startswith('-')]
        items.append({'company_name': parts[0], 'duration': parts[1], 'location': parts[2], 'bullets': bullets})
    return items


def _serialize_education_history(items: list[dict]) -> str:
    return '\n\n'.join(' | '.join([item.get('university', ''), item.get('degree', ''), item.get('duration', ''), item.get('location', '')]).strip() for item in items)


def _parse_education_history(value: str) -> list[dict]:
    items: list[dict] = []
    for block in [part.strip() for part in str(value or '').split('\n\n') if part.strip()]:
        parts = [part.strip() for part in block.split('|')]
        while len(parts) < 4:
            parts.append('')
        items.append({'university': parts[0], 'degree': parts[1], 'duration': parts[2], 'location': parts[3]})
    return items


def _serialize_skill_groups(groups: list[dict]) -> str:
    lines: list[str] = []
    for group in groups or []:
        category = str(group.get('category', '')).strip()
        items = ', '.join(str(item).strip() for item in group.get('items', []) if str(item).strip())
        if category and items:
            lines.append(f'{category}: {items}')
    return '\n'.join(lines)


def _parse_skill_groups(value: str) -> list[dict]:
    groups: list[dict] = []
    for line in [item.strip() for item in str(value or '').splitlines() if item.strip()]:
        if ':' in line:
            category, raw_items = line.split(':', 1)
        else:
            category, raw_items = 'Other Relevant', line
        items = _parse_comma_separated_list(raw_items)
        if category.strip() and items:
            groups.append({'category': category.strip(), 'items': items})
    return groups


def _queue_editor_reload(resume: dict, notice: str = '') -> None:
    st.session_state['editor_pending_resume'] = resume or {}
    st.session_state['editor_notice'] = notice


def _load_editor_from_resume(resume: dict, force: bool = False) -> None:
    signature = json.dumps(resume or {}, sort_keys=True, ensure_ascii=False)
    if not force and st.session_state.get('editor_loaded_signature') == signature:
        return
    resume = resume or {}
    st.session_state['editor_headline'] = resume.get('headline', '')
    st.session_state['editor_summary'] = resume.get('summary', '')
    st.session_state['editor_technical_skills'] = ', '.join(resume.get('technical_skills', []))
    st.session_state['editor_fit_keywords'] = ', '.join(resume.get('fit_keywords', []))
    st.session_state['editor_skill_groups'] = _serialize_skill_groups(resume.get('skill_groups', []))
    st.session_state['editor_education'] = _serialize_education_history(resume.get('education_history', []))
    st.session_state['editor_work_count'] = len(resume.get('work_history', []))
    for idx, job in enumerate(resume.get('work_history', [])):
        st.session_state[f'editor_job_company_{idx}'] = job.get('company_name', '')
        st.session_state[f'editor_job_duration_{idx}'] = job.get('duration', '')
        st.session_state[f'editor_job_location_{idx}'] = job.get('location', '')
        st.session_state[f'editor_job_role_{idx}'] = job.get('role_title', '')
        st.session_state[f'editor_job_headline_{idx}'] = job.get('role_headline', '')
        st.session_state[f'editor_job_bullets_{idx}'] = '\n'.join(job.get('bullets', []))
    st.session_state['editor_loaded_signature'] = signature


def _resume_from_editor(base_resume: dict) -> dict:
    skill_groups = _parse_skill_groups(st.session_state.get('editor_skill_groups', ''))
    technical_skills = _parse_comma_separated_list(st.session_state.get('editor_technical_skills', ''))
    if skill_groups:
        grouped_items: list[str] = []
        for group in skill_groups:
            grouped_items.extend(group.get('items', []))
        technical_skills = _dedupe_preserve_order(technical_skills + grouped_items)
    work_history: list[dict] = []
    for idx in range(int(st.session_state.get('editor_work_count', 0))):
        bullets = [line.strip() for line in st.session_state.get(f'editor_job_bullets_{idx}', '').splitlines() if line.strip()]
        work_history.append({'company_name': st.session_state.get(f'editor_job_company_{idx}', '').strip(), 'duration': st.session_state.get(f'editor_job_duration_{idx}', '').strip(), 'location': st.session_state.get(f'editor_job_location_{idx}', '').strip(), 'role_title': st.session_state.get(f'editor_job_role_{idx}', '').strip(), 'role_headline': st.session_state.get(f'editor_job_headline_{idx}', '').strip(), 'bullets': bullets})
    return {'headline': st.session_state.get('editor_headline', '').strip(), 'summary': st.session_state.get('editor_summary', '').strip(), 'technical_skills': technical_skills, 'skill_groups': skill_groups, 'fit_keywords': _parse_comma_separated_list(st.session_state.get('editor_fit_keywords', '')), 'work_history': work_history, 'education_history': _parse_education_history(st.session_state.get('editor_education', '')), 'bold_keywords': base_resume.get('bold_keywords', []), 'auto_bold_fit_keywords': bool(base_resume.get('auto_bold_fit_keywords', False))}


# ---------- Navigation ----------

def render_top_nav(user: dict) -> str:
    options = ['Dashboard', 'Job List', 'Generated Resumes']
    if is_admin(user):
        options += ['User Access', 'Profile Settings', 'App Settings']
    else:
        options += ['My Weekly Result']

    latest_resume = st.session_state.get('last_resume') or {}
    latest_jd = st.session_state.get('last_job_description', '')
    current_score = None
    if latest_resume and str(latest_jd).strip():
        try:
            current_score = analyze_ats_score(latest_resume, latest_jd, target_role=st.session_state.get('last_target_role', '')).get('overall_score', 0)
        except Exception:
            current_score = None

    pending_nav = st.session_state.pop('pending_nav_page', '')
    if pending_nav and pending_nav in options:
        st.session_state['top_navigation'] = pending_nav
    elif st.session_state.get('top_navigation') not in options:
        st.session_state['top_navigation'] = options[0]

    current_page = st.session_state.get('top_navigation', options[0])

    generated_total = len(storage.get_generated_resumes())
    if not is_admin(user):
        generated_total = len([item for item in storage.get_generated_resumes() if item.get('created_by_user_id') == user.get('id')])

    def _label(page: str) -> str:
        if page == 'Dashboard' and current_score is not None:
            return f'{page} ({current_score})'
        if page == 'Generated Resumes':
            return f'{page} ({generated_total})'
        return page

    st.markdown(
        """
        <style>
        .main > div.block-container {padding-top: 0.9rem; padding-bottom: 1.5rem;}
        .top-user-meta {text-align:right; margin-top: 0.1rem;}
        .top-user-name {font-weight:700; font-size:0.96rem; margin:0;}
        .top-user-role {color:#94a3b8; font-size:0.82rem; margin:0.08rem 0 0 0;}
        div[data-testid="stButton"] > button {border-radius:12px; min-height:2.7rem;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    nav_shell, user_shell = st.columns([5.8, 1.8], gap='medium')
    with nav_shell:
        nav_cols = st.columns(len(options), gap='small')
        for idx, page in enumerate(options):
            with nav_cols[idx]:
                if st.button(_label(page), key=f'nav_btn_{page}', use_container_width=True, type='primary' if current_page == page else 'secondary'):
                    st.session_state['top_navigation'] = page
                    st.rerun()
    with user_shell:
        meta_col, logout_col = st.columns([1.25, 1.0], gap='small')
        with meta_col:
            st.markdown(
                f"<div class='top-user-meta'><p class='top-user-name'>{html.escape(user.get('full_name') or user.get('username') or '')}</p><p class='top-user-role'>{'Admin' if is_admin(user) else 'User'}</p></div>",
                unsafe_allow_html=True,
            )
        with logout_col:
            if st.button('Logout', key='top_logout_btn', use_container_width=True):
                _clear_login_token()
                st.session_state['current_user_id'] = ''
                st.session_state['pending_nav_page'] = ''
                st.session_state['auth_notice'] = 'You are signed out.'
                st.rerun()

    return current_page


# ---------- Boot ----------

init_state()
_restore_auth_from_token()
start_job_scrape_worker(str(APP_DIR / 'data'))
current_user = require_auth()
page = render_top_nav(current_user)

if page == 'Dashboard':
    dashboard_page(current_user)
elif page == 'Job List':
    job_list_page(current_user)
elif page == 'Generated Resumes':
    generated_resumes_page(current_user)
elif page == 'My Weekly Result':
    my_weekly_result_page(current_user)
elif page == 'User Access':
    user_access_page(current_user)
elif page == 'Profile Settings':
    profile_settings_page(current_user)
else:
    app_settings_page(current_user)
