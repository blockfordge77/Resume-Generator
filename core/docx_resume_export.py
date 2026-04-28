from __future__ import annotations

import base64
import html
import json
import re
import shutil
import subprocess
import sys
import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


SECTION_ALIASES = {
    "summary": {
        "summary",
        "professional summary",
        "profile",
        "about",
        "about me",
        "professional profile",
        "profile summary",
        "career profile",
        "career summary",
    },
    "skills": {
        "skills",
        "technical skills",
        "core skills",
        "core competencies",
        "technologies",
        "technology stack",
    },
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment history",
        "work history",
    },
    "education": {
        "education",
        "education history",
        "academic background",
    },
    "projects": {"projects", "selected projects"},
    "certifications": {"certifications", "certificates", "licenses"},
}
ALL_SECTION_TITLES = {title for titles in SECTION_ALIASES.values() for title in titles}
TITLE_PLACEHOLDERS = {"___resume_title___", "___headline___", "__resume_title__", "__headline__"}
EXP_ROLE_PLACEHOLDERS = {"___title___", "__role__"}
SUMMARY_PLACEHOLDERS = {"___summary___", "___professional_summary___"}
SKILL_PLACEHOLDERS = {"___skills___", "___technical_skills___"}
EXPERIENCE_PLACEHOLDERS = {"___experience___", "___work_experience___", "___professional_experience___"}

_DATE_RE = re.compile(
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}\b|\b\d{4}\s*[–—-]\s*(?:present|current|\d{4})\b",
    re.IGNORECASE,
)
_ROLE_MARKER_PATTERN = re.compile(r"(___title___|__role__)", re.IGNORECASE)


def _clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _strip_markdown_emphasis(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(?![a-z_]+__)(.*?)__", r"\1", text)
    return text


def _plain_resume_text(value: object) -> str:
    return _clean_text(_strip_markdown_emphasis(value))


def _preserve_line_text(value: object) -> str:
    return _strip_markdown_emphasis(value).replace("\r", "").replace("\n", " ")


def _paragraph_contains_any_marker(paragraph: Paragraph, markers: set[str]) -> bool:
    text_lower = (paragraph.text or "").lower()
    return any(marker.lower() in text_lower for marker in markers)


def _paragraphs_in_container(container) -> Iterable[Paragraph]:
    """Yield paragraphs in real visual DOCX order, including tables.

    Some resumes use one-cell tables only for section headers. Using
    container.paragraphs first and container.tables second makes section ranges
    wrong. Walking the OOXML children preserves the visible order:
    paragraph -> table heading -> following outside-table content.
    """
    if hasattr(container, "_body"):
        parent_elm = getattr(container._body, "_element", None)
    else:
        parent_elm = getattr(container, "_element", None) or getattr(container, "element", None)

    if parent_elm is None:
        for paragraph in getattr(container, "paragraphs", []) or []:
            yield paragraph
        for table in getattr(container, "tables", []) or []:
            for row in table.rows:
                for cell in row.cells:
                    yield from _paragraphs_in_container(cell)
        return

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, container)
        elif child.tag == qn("w:tbl"):
            table = Table(child, container)
            seen_cells: set[int] = set()
            for row in table.rows:
                for cell in row.cells:
                    key = id(cell._tc)
                    if key in seen_cells:
                        continue
                    seen_cells.add(key)
                    yield from _paragraphs_in_container(cell)


def _all_story_paragraphs(doc: Document) -> list[Paragraph]:
    paragraphs = list(_paragraphs_in_container(doc))
    seen_parts = {id(doc.part)}
    for section in doc.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            try:
                part_id = id(container.part)
            except Exception:
                continue
            if part_id in seen_parts:
                continue
            seen_parts.add(part_id)
            paragraphs.extend(list(_paragraphs_in_container(container)))
    return paragraphs


def _all_body_paragraphs(doc: Document) -> list[Paragraph]:
    return list(_paragraphs_in_container(doc))


def _first_run_or_none(paragraph: Paragraph):
    return paragraph.runs[0] if paragraph.runs else None


def _first_meaningful_run_or_none(paragraph: Paragraph):
    for run in paragraph.runs:
        if str(run.text or "").strip():
            return run
    return _first_run_or_none(paragraph)


def _clear_paragraph_keep_ppr(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _copy_run_format(source_run, target_run) -> None:
    if source_run is None:
        return
    try:
        if source_run._r.rPr is not None:
            target_run._r.insert(0, deepcopy(source_run._r.rPr))
    except Exception:
        pass


def _force_run_not_bold(run) -> None:
    try:
        run.bold = False
    except Exception:
        pass
    try:
        rpr = run._r.get_or_add_rPr()
        for tag in ("w:b", "w:bCs"):
            node = rpr.find(qn(tag))
            if node is not None:
                rpr.remove(node)
        rstyle = rpr.find(qn("w:rStyle"))
        if rstyle is not None:
            style_id = str(rstyle.get(qn("w:val"), "")).strip().lower()
            if style_id in {"strong", "bold", "13"}:
                rpr.remove(rstyle)
    except Exception:
        pass


def _force_run_bold(run) -> None:
    try:
        run.bold = True
    except Exception:
        pass
    try:
        rpr = run._r.get_or_add_rPr()
        if rpr.find(qn("w:b")) is None:
            rpr.append(OxmlElement("w:b"))
        if rpr.find(qn("w:bCs")) is None:
            rpr.append(OxmlElement("w:bCs"))
    except Exception:
        pass


def _technical_skill_keywords(resume: dict) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()
    for item in resume.get("technical_skills", []) or []:
        clean = _plain_resume_text(item)
        if len(clean) < 2:
            continue
        key = clean.lower()
        if key not in seen:
            seen.add(key)
            ordered.append(clean)
    return ordered


def _keyword_pattern(keywords: list[str]) -> re.Pattern[str]:
    clean_keywords = sorted(
        {str(item).strip() for item in keywords if str(item).strip()},
        key=len,
        reverse=True,
    )
    if not clean_keywords:
        return re.compile(r"(?!x)x")
    parts: list[str] = []
    for keyword in clean_keywords:
        escaped = re.escape(keyword)
        # Avoid bolding inside larger words, while allowing punctuation in tech names.
        parts.append(rf"(?<![A-Za-z0-9]){escaped}(?![A-Za-z0-9])")
    return re.compile("|".join(parts), re.IGNORECASE)


def _add_run(
    paragraph: Paragraph,
    text: str,
    source_run=None,
    *,
    force_no_bold: bool = False,
    force_bold: bool = False,
):
    run = paragraph.add_run(text)
    _copy_run_format(source_run, run)
    if force_no_bold:
        _force_run_not_bold(run)
    if force_bold:
        _force_run_bold(run)
    return run


def _add_text_with_keyword_bold(
    paragraph: Paragraph,
    text: str,
    source_run=None,
    keywords: list[str] | None = None,
    *,
    base_no_bold: bool = True,
) -> None:
    text = _plain_resume_text(text)
    keywords = keywords or []
    if not keywords:
        _add_run(paragraph, text, source_run, force_no_bold=base_no_bold)
        return
    pattern = _keyword_pattern(keywords)
    pos = 0
    matched_any = False
    for match in pattern.finditer(text):
        if match.start() > pos:
            _add_run(paragraph, text[pos:match.start()], source_run, force_no_bold=base_no_bold)
        _add_run(paragraph, match.group(0), source_run, force_no_bold=True, force_bold=True)
        matched_any = True
        pos = match.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], source_run, force_no_bold=base_no_bold)
    if not matched_any and not text:
        _add_run(paragraph, "", source_run, force_no_bold=base_no_bold)


def _set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    source_run=None,
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> None:
    source_run = source_run if source_run is not None else _first_run_or_none(paragraph)
    _clear_paragraph_keep_ppr(paragraph)
    if keywords:
        _add_text_with_keyword_bold(paragraph, text, source_run, keywords, base_no_bold=True)
    else:
        _add_run(paragraph, _plain_resume_text(text), source_run, force_no_bold=force_no_bold)


def _insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    like_paragraph: Paragraph | None = None,
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> Paragraph:
    like_paragraph = like_paragraph or paragraph
    new_p = OxmlElement("w:p")
    if like_paragraph._p.pPr is not None:
        new_p.append(deepcopy(like_paragraph._p.pPr))
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    source_run = _first_meaningful_run_or_none(like_paragraph)
    if keywords:
        _add_text_with_keyword_bold(new_paragraph, text, source_run, keywords, base_no_bold=True)
    else:
        _add_run(new_paragraph, _plain_resume_text(text), source_run, force_no_bold=force_no_bold)
    return new_paragraph


def _delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def _is_decorative_or_blank_paragraph(paragraph: Paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return True
    return bool(re.fullmatch(r"[_\-—–=]{5,}", text))


def _replace_paragraph_with_lines(
    paragraph: Paragraph,
    lines: list[str],
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> None:
    lines = [str(line or "").strip() for line in lines if str(line or "").strip()]
    if not lines:
        _set_paragraph_text(paragraph, "", force_no_bold=force_no_bold)
        return
    source_run = _first_meaningful_run_or_none(paragraph)
    _set_paragraph_text(paragraph, lines[0], source_run, keywords=keywords, force_no_bold=force_no_bold)
    cursor = paragraph
    for line in lines[1:]:
        cursor = _insert_paragraph_after(cursor, line, like_paragraph=paragraph, keywords=keywords, force_no_bold=force_no_bold)


def _normalized_heading_text(paragraph: Paragraph) -> str:
    return _clean_text(paragraph.text).lower().strip(":")


def _is_section_heading(paragraph: Paragraph) -> bool:
    text = _normalized_heading_text(paragraph)
    if text in ALL_SECTION_TITLES:
        return True
    style_name = str(getattr(paragraph.style, "name", "") or "").lower()
    return bool(text) and "heading" in style_name and len(text.split()) <= 4


def _section_name(paragraph: Paragraph) -> str:
    text = _normalized_heading_text(paragraph)
    for key, aliases in SECTION_ALIASES.items():
        if text in aliases:
            return key
    return ""


def _find_section_range(paragraphs: list[Paragraph], section_key: str) -> tuple[int, int] | None:
    start = -1
    for idx, paragraph in enumerate(paragraphs):
        if _section_name(paragraph) == section_key:
            start = idx
            break
    if start < 0:
        return None
    end = len(paragraphs)
    for idx in range(start + 1, len(paragraphs)):
        if _is_section_heading(paragraphs[idx]):
            end = idx
            break
    return start, end


def _paragraphs_between_sections(doc: Document, section_key: str) -> list[Paragraph]:
    paragraphs = _all_body_paragraphs(doc)
    section_range = _find_section_range(paragraphs, section_key)
    if not section_range:
        return []
    start, end = section_range
    return paragraphs[start + 1:end]


def _replace_section_body(
    doc: Document,
    section_key: str,
    lines: list[str],
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> bool:
    paragraphs = _all_body_paragraphs(doc)
    section_range = _find_section_range(paragraphs, section_key)
    if not section_range:
        return False
    start, end = section_range
    body = paragraphs[start + 1:end]
    content_body = [p for p in body if _clean_text(p.text) and not _is_decorative_or_blank_paragraph(p)]
    if content_body:
        anchor = content_body[0]
        # Preserve blank/decorative paragraphs; remove only old content body.
        for p in content_body[1:]:
            _delete_paragraph(p)
        _replace_paragraph_with_lines(anchor, lines, keywords=keywords, force_no_bold=force_no_bold)
    else:
        anchor = paragraphs[start]
        cursor = anchor
        for line in [line for line in lines if str(line).strip()]:
            cursor = _insert_paragraph_after(cursor, line, like_paragraph=anchor, keywords=keywords, force_no_bold=force_no_bold)
    return True


def _replace_placeholders(
    doc: Document,
    placeholders: set[str],
    lines: list[str],
    *,
    keywords: list[str] | None = None,
    force_no_bold: bool = False,
) -> bool:
    changed = False
    lowered = {item.lower() for item in placeholders}
    for paragraph in _all_story_paragraphs(doc):
        text = _clean_text(paragraph.text).lower()
        if text in lowered:
            _replace_paragraph_with_lines(paragraph, lines, keywords=keywords, force_no_bold=force_no_bold)
            changed = True
    return changed


def _copy_paragraph_text_with_replacements(paragraph: Paragraph, replacements: dict[str, str]) -> bool:
    if not replacements:
        return False
    cleaned = {marker: _plain_resume_text(value) for marker, value in replacements.items()}
    original = paragraph.text or ""
    desired = original
    for marker, value in cleaned.items():
        desired = re.sub(re.escape(marker), value, desired, flags=re.IGNORECASE)
    if desired == original:
        return False

    changed_in_runs = False
    for run in paragraph.runs:
        run_text = run.text or ""
        new_text = run_text
        for marker, value in cleaned.items():
            new_text = re.sub(re.escape(marker), value, new_text, flags=re.IGNORECASE)
        if new_text != run_text:
            run.text = _preserve_line_text(new_text)
            changed_in_runs = True
    if changed_in_runs and (paragraph.text or "") == desired:
        return True

    _set_paragraph_text(paragraph, desired, _first_meaningful_run_or_none(paragraph))
    return True


def _replace_inline_placeholders(doc: Document, replacements: dict[str, str]) -> bool:
    changed = False
    for paragraph in _all_story_paragraphs(doc):
        if _copy_paragraph_text_with_replacements(paragraph, replacements):
            changed = True
    return changed


def _set_keep_lines(paragraph: Paragraph) -> None:
    try:
        paragraph.paragraph_format.keep_together = True
    except Exception:
        pass
    try:
        ppr = paragraph._p.get_or_add_pPr()
        if ppr.find(qn("w:keepLines")) is None:
            ppr.append(OxmlElement("w:keepLines"))
    except Exception:
        pass


def _remove_existing_tabs(ppr) -> None:
    try:
        tabs = ppr.find(qn("w:tabs"))
        if tabs is not None:
            ppr.remove(tabs)
    except Exception:
        pass


def _content_width_twips(paragraph: Paragraph) -> int:
    try:
        section = paragraph.part.document.sections[0]
        return int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)
    except Exception:
        return 9000


def _set_role_line_right_tab(paragraph: Paragraph) -> None:
    """Align right-side date/meta with a right tab instead of many spaces.

    Long generated roles cannot be kept stable with raw spaces in DOCX/PDF.
    A right tab preserves the line visually and prevents WPS/Word from expanding
    100+ spaces during PDF export.
    """
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass
    try:
        ppr = paragraph._p.get_or_add_pPr()
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            ppr.remove(jc)
        _remove_existing_tabs(ppr)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(_content_width_twips(paragraph)))
        tabs.append(tab)
        ppr.append(tabs)
    except Exception:
        pass


def _append_tab(paragraph: Paragraph, source_run=None) -> None:
    run = paragraph.add_run()
    _copy_run_format(source_run, run)
    try:
        run.add_tab()
    except Exception:
        run.text = "\t"


def _split_role_line(original: str):
    text = str(original or "").replace("\r", "").replace("\n", " ")
    match = _ROLE_MARKER_PATTERN.search(text)
    if not match:
        return None
    before = text[:match.start()]
    after_marker = text[match.end():]
    pipe_offset = after_marker.find("|")
    if pipe_offset < 0:
        return {
            "before": before,
            "role_marker": match.group(0),
            "separator": "",
            "right_text": after_marker,
            "after": "",
            "has_pipe": False,
        }
    bounded = after_marker[:pipe_offset]
    after = after_marker[pipe_offset + 1:]
    sep_match = re.match(r"([ \t\u00a0]*)(.*)", bounded, flags=re.DOTALL)
    separator = sep_match.group(1) if sep_match else ""
    right_text = (sep_match.group(2) if sep_match else bounded).strip()
    return {
        "before": before,
        "role_marker": match.group(0),
        "separator": separator,
        "right_text": right_text,
        "after": after,
        "has_pipe": True,
    }


def _replace_role_paragraph(paragraph: Paragraph, role_value: str) -> bool:
    original = paragraph.text or ""
    parts = _split_role_line(original)
    if not parts:
        return False
    role = _plain_resume_text(role_value)
    if not role:
        return False

    if not parts["has_pipe"]:
        replacements = {marker: role for marker in EXP_ROLE_PLACEHOLDERS}
        return _copy_paragraph_text_with_replacements(paragraph, replacements)

    before_run = _first_meaningful_run_or_none(paragraph)
    role_run = before_run
    right_run = before_run
    _clear_paragraph_keep_ppr(paragraph)
    _set_role_line_right_tab(paragraph)
    if parts["before"]:
        _add_run(paragraph, str(parts["before"]), before_run)
    _add_run(paragraph, role, role_run)
    if parts["right_text"]:
        _append_tab(paragraph, role_run)
        _add_run(paragraph, str(parts["right_text"]), right_run)
    if parts["after"]:
        _add_run(paragraph, str(parts["after"]).replace("|", ""), right_run)
    _set_keep_lines(paragraph)
    return True


def _replace_role_placeholders(doc: Document, resume: dict) -> bool:
    jobs = resume.get("work_history", []) or []
    role_values = [_plain_resume_text(job.get("role_title", "")) for job in jobs]
    role_values = [value for value in role_values if value]
    fallback_role = role_values[0] if role_values else _plain_resume_text(resume.get("headline", ""))
    if not fallback_role:
        return False
    role_index = 0
    changed = False
    for paragraph in _all_story_paragraphs(doc):
        if not _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
            continue
        role_value = role_values[min(role_index, len(role_values) - 1)] if role_values else fallback_role
        role_index += 1
        if _replace_role_paragraph(paragraph, role_value):
            changed = True
    return changed


def _parse_skill_groups(value) -> list[dict]:
    if isinstance(value, str):
        raw = value.strip()
        if raw:
            try:
                parsed = json.loads(raw)
                if isinstance(parsed, list):
                    value = parsed
            except Exception:
                value = []
    return value if isinstance(value, list) else []


def _skill_lines(resume: dict) -> list[str]:
    groups = _parse_skill_groups(resume.get("skill_groups") or [])
    lines: list[str] = []
    for group in groups:
        if not isinstance(group, dict):
            continue
        category = _plain_resume_text(group.get("category", ""))
        items = [_plain_resume_text(item) for item in group.get("items", []) if _plain_resume_text(item)]
        if category and items:
            lines.append(f"{category}: {', '.join(items)}")
    if lines:
        return lines
    skills = [_plain_resume_text(item) for item in resume.get("technical_skills", []) if _plain_resume_text(item)]
    return [", ".join(skills)] if skills else []


def _paragraph_has_tab(paragraph: Paragraph) -> bool:
    return "\t" in (paragraph.text or "")


def _paragraph_has_numbering(paragraph: Paragraph) -> bool:
    try:
        ppr = paragraph._p.pPr
        return bool(ppr is not None and ppr.numPr is not None)
    except Exception:
        return False


def _looks_like_job_meta_line(paragraph: Paragraph) -> bool:
    """Protect company/location/duration rows from bullet replacement.

    Marvin-style resumes sometimes store company/location/duration as a
    numbered paragraph, so a pure numbering test misclassifies it as a bullet.
    These rows usually contain a company/location separator and a date range,
    and are often bold. They must stay unchanged unless explicit placeholders
    are present.
    """
    text = _clean_text(paragraph.text)
    if not text:
        return False
    if _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
        return False
    has_date = bool(_DATE_RE.search(text))
    has_company_separator = " | " in text or "|" in text
    has_many_spaces = bool(re.search(r"\S\s{8,}\S", paragraph.text or ""))
    if has_date and (has_company_separator or has_many_spaces):
        return True
    # Strong fallback for rows such as "Company Location  Jan 2020 – Present".
    if has_date and len(text.split()) <= 14 and not text.endswith("."):
        return True
    return False


def _looks_like_role_title_line(paragraph: Paragraph) -> bool:
    text = _clean_text(paragraph.text)
    if not text:
        return False
    if _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
        return False
    if _DATE_RE.search(text) or "|" in text:
        return False
    if text.endswith("."):
        return False
    words = text.split()
    if len(words) > 8:
        return False
    role_words = {"engineer", "developer", "lead", "manager", "architect", "analyst", "consultant", "specialist", "intern", "director"}
    return any(word.strip("(),/-").lower() in role_words for word in words)


def _paragraph_looks_like_bullet_content(paragraph: Paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return False
    if _paragraph_contains_any_marker(paragraph, EXP_ROLE_PLACEHOLDERS):
        return False
    if _looks_like_job_meta_line(paragraph) or _looks_like_role_title_line(paragraph):
        return False
    if _paragraph_has_tab(paragraph):
        return False
    if text.endswith(":") and len(text.split()) <= 5:
        return False
    style_name = str(getattr(paragraph.style, "name", "") or "").lower()
    if "heading" in style_name:
        return False
    if text.startswith(("•", "-", "*", "‣", "◦")):
        return True
    if "bullet" in style_name:
        return True
    if "list" in style_name and len(text.split()) >= 6 and text.endswith((".", ";")):
        return True
    if _paragraph_has_numbering(paragraph) and len(text.split()) >= 6 and text.endswith((".", ";")):
        return True
    return False


def _replace_experience_bullets_and_titles(doc: Document, resume: dict) -> bool:
    body = _paragraphs_between_sections(doc, "experience")
    if not body:
        return False
    jobs = resume.get("work_history", []) or []
    if not jobs:
        return False
    changed = False
    keywords = _technical_skill_keywords(resume)

    bullet_groups: list[list[Paragraph]] = []
    current_group: list[Paragraph] = []
    for paragraph in body:
        if _paragraph_looks_like_bullet_content(paragraph):
            current_group.append(paragraph)
        else:
            if current_group:
                bullet_groups.append(current_group)
                current_group = []
    if current_group:
        bullet_groups.append(current_group)

    for idx, group in enumerate(bullet_groups):
        if idx >= len(jobs) or not group:
            continue
        bullets = [_plain_resume_text(bullet) for bullet in jobs[idx].get("bullets", []) if _plain_resume_text(bullet)]
        if not bullets:
            continue

        first = group[0]
        # Delete extra old bullets only inside this real bullet group; company,
        # role, duration, and location rows are not in this group.
        for old_paragraph in group[len(bullets):]:
            _delete_paragraph(old_paragraph)

        active_group = group[:len(bullets)]
        for paragraph, bullet in zip(active_group, bullets[:len(active_group)]):
            _set_paragraph_text(paragraph, bullet, keywords=keywords, force_no_bold=True)
            changed = True

        cursor = active_group[-1] if active_group else first
        for bullet in bullets[len(active_group):]:
            cursor = _insert_paragraph_after(cursor, bullet, like_paragraph=first, keywords=keywords, force_no_bold=True)
            changed = True

    return changed


def apply_resume_to_docx(docx_path: Path, resume: dict) -> None:
    doc = Document(str(docx_path))
    headline = _plain_resume_text(resume.get("headline", ""))
    summary = _plain_resume_text(resume.get("summary", ""))
    skills = _skill_lines(resume)
    tech_keywords = _technical_skill_keywords(resume)

    if headline:
        _replace_inline_placeholders(
            doc,
            {
                "___resume_title___": headline,
                "___headline___": headline,
                "__resume_title__": headline,
                "__headline__": headline,
            },
        )
    _replace_role_placeholders(doc, resume)

    if summary:
        if not _replace_placeholders(doc, SUMMARY_PLACEHOLDERS, [summary], keywords=tech_keywords, force_no_bold=True):
            _replace_section_body(doc, "summary", [summary], keywords=tech_keywords, force_no_bold=True)

    if skills:
        # Skills section must always be normal/no-bold, even when terms are in
        # technical_skills or the uploaded sample skills paragraph is bold.
        if not _replace_placeholders(doc, SKILL_PLACEHOLDERS, skills, force_no_bold=True):
            _replace_section_body(doc, "skills", skills, force_no_bold=True)

    # Experience placeholders replace an entire placeholder block. Normal
    # experience sections update only bullet paragraphs to protect company,
    # role, location, and duration style.
    experience_lines = []
    for job_index, job in enumerate(resume.get("work_history", []) or []):
        if job_index:
            experience_lines.append("")
        meta = " | ".join(
            item
            for item in [
                _plain_resume_text(job.get("company_name", "")),
                _plain_resume_text(job.get("duration", "")),
                _plain_resume_text(job.get("location", "")),
            ]
            if item
        )
        if meta:
            experience_lines.append(meta)
        role = _plain_resume_text(job.get("role_title", ""))
        if role:
            experience_lines.append(role)
        for bullet in job.get("bullets", []) or []:
            bullet_text = _plain_resume_text(bullet)
            if bullet_text:
                experience_lines.append(f"• {bullet_text}")

    if experience_lines:
        if not _replace_placeholders(doc, EXPERIENCE_PLACEHOLDERS, experience_lines, keywords=tech_keywords, force_no_bold=True):
            _replace_experience_bullets_and_titles(doc, resume)

    doc.save(str(docx_path))


def find_soffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("soffice.exe"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/snap/bin/libreoffice",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def find_wps() -> str | None:
    candidates = [
        shutil.which("wps"),
        shutil.which("wps.exe"),
        r"C:\Program Files\WPS Office\office6\wps.exe",
        r"C:\Program Files (x86)\WPS Office\office6\wps.exe",
        r"C:\Program Files\Kingsoft\WPS Office\office6\wps.exe",
        r"C:\Program Files (x86)\Kingsoft\WPS Office\office6\wps.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def export_pdf_via_docx2pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    try:
        from docx2pdf import convert  # type: ignore
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "PDF created via docx2pdf"
        return False, "docx2pdf ran but no PDF file was created"
    except Exception as exc:
        return False, f"docx2pdf failed: {exc!r}"


def export_pdf_via_word(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    if sys.platform != "win32":
        return False, "Word COM export is only supported on Windows"
    word = None
    doc = None
    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore
        pythoncom.CoInitialize()
        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass
        doc = word.Documents.Open(str(docx_path), ReadOnly=True, AddToRecentFiles=False, ConfirmConversions=False)
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            CreateBookmarks=1,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "PDF created via Microsoft Word"
        return False, "Word export ran but no PDF file was created"
    except Exception as exc:
        return False, f"Word export failed: {exc!r}"
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            import pythoncom  # type: ignore
            pythoncom.CoUninitialize()
        except Exception:
            pass


def export_pdf_via_libreoffice(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    soffice = find_soffice()
    if not soffice:
        return False, "LibreOffice not found"
    temp_dir = Path(tempfile.mkdtemp(prefix="lo_pdf_"))
    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(temp_dir), str(docx_path)],
            capture_output=True,
            text=True,
            check=False,
            timeout=90,
        )
        generated = temp_dir / f"{docx_path.stem}.pdf"
        if generated.exists() and generated.stat().st_size > 0:
            pdf_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(generated, pdf_path)
            return True, "PDF created via LibreOffice"
        message = result.stderr.strip() or result.stdout.strip() or "LibreOffice conversion failed"
        return False, f"LibreOffice export failed: {message}"
    except Exception as exc:
        return False, f"LibreOffice export failed: {exc!r}"
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def export_pdf_via_wps_custom(docx_path: Path, pdf_path: Path, pdf_cfg: dict) -> tuple[bool, str]:
    command_template = str(pdf_cfg.get("wps_pdf_command", "") or "").strip()
    if not command_template:
        if find_wps():
            return False, "WPS found, but no WPS custom PDF command is configured"
        return False, "WPS not found and no WPS custom PDF command is configured"
    command = command_template.format(input=str(docx_path), output=str(pdf_path))
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True, check=False)
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "PDF created via WPS custom command"
        message = result.stderr.strip() or result.stdout.strip() or "WPS custom command failed"
        return False, f"WPS export failed: {message}"
    except Exception as exc:
        return False, f"WPS export failed: {exc!r}"


def export_pdf(docx_path: Path, pdf_path: Path, pdf_cfg: dict | None = None) -> tuple[bool, str]:
    pdf_cfg = pdf_cfg or {}
    order = pdf_cfg.get("backend_order")
    if isinstance(order, str):
        order = [item.strip() for item in order.split(",") if item.strip()]
    if not isinstance(order, list) or not order:
        order = ["docx2pdf", "word", "libreoffice", "wps_custom"]
    backend_map = {
        "docx2pdf": lambda: export_pdf_via_docx2pdf(docx_path, pdf_path),
        "word": lambda: export_pdf_via_word(docx_path, pdf_path),
        "libreoffice": lambda: export_pdf_via_libreoffice(docx_path, pdf_path),
        "wps_custom": lambda: export_pdf_via_wps_custom(docx_path, pdf_path, pdf_cfg),
    }
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    messages: list[str] = []
    for backend in order:
        key = str(backend or "").strip().lower()
        fn = backend_map.get(key)
        if not fn:
            messages.append(f"{backend}: unknown backend")
            continue
        ok, message = fn()
        messages.append(f"{backend}: {message}")
        if ok:
            return True, message
    return False, " | ".join(messages)


def pdf_backend_status(pdf_cfg: dict | None = None) -> list[str]:
    pdf_cfg = pdf_cfg or {}
    lines = []
    try:
        from docx2pdf import convert  # noqa: F401
        lines.append("docx2pdf: OK - package available")
    except Exception as exc:
        lines.append(f"docx2pdf: NO - {exc!r}")
    lines.append("word: OK - Windows only; checked during export" if sys.platform == "win32" else "word: NO - not Windows")
    soffice = find_soffice()
    lines.append(f"libreoffice: OK - {soffice}" if soffice else "libreoffice: NO - not found")
    wps = find_wps()
    if str(pdf_cfg.get("wps_pdf_command", "") or "").strip():
        lines.append("wps_custom: OK - command configured")
    elif wps:
        lines.append(f"wps_custom: NO - WPS found at {wps}, but command is not configured")
    else:
        lines.append("wps_custom: NO - WPS not found")
    return lines


def _uploaded_resume_path(profile: dict) -> Path | None:
    upload = profile.get("uploaded_resume") if isinstance(profile.get("uploaded_resume"), dict) else {}
    candidates = [
        str(upload.get("path", "") or "").strip(),
        str(upload.get("storage_path", "") or "").strip(),
        str(upload.get("relative_path", "") or "").strip(),
    ]
    for value in candidates:
        if not value:
            continue
        path = Path(value).expanduser()
        if path.exists():
            return path
    return None


def build_pdf_preview_html(pdf_bytes: bytes, message: str = "") -> str:
    if not pdf_bytes:
        return f"""
        <div style='font-family:Arial,sans-serif;padding:20px;border:1px solid #fca5a5;border-radius:12px;background:#fff1f2;color:#7f1d1d;'>
          <h3 style='margin-top:0;'>PDF preview is unavailable</h3>
          <p>{html.escape(message or 'The PDF exporter did not return a PDF file.')}</p>
        </div>
        """
    encoded = base64.b64encode(pdf_bytes).decode("ascii")
    return f"""
    <div style='font-family:Arial,sans-serif;'>
      <iframe title='Resume PDF preview' src='data:application/pdf;base64,{encoded}' style='width:100%;height:1120px;border:1px solid #e5e7eb;border-radius:12px;background:#fff;'></iframe>
      <p style='font-size:12px;color:#64748b;margin-top:8px;'>Read-only PDF preview generated from the uploaded DOCX style. {html.escape(message)}</p>
    </div>
    """


def build_docx_style_pdf_bundle(resume: dict, profile: dict, output_dir: Path | str, pdf_cfg: dict | None = None) -> dict[str, bytes | str]:
    source_docx = _uploaded_resume_path(profile)
    if not source_docx:
        raise FileNotFoundError("no resume so must upload resume")
    temp_dir = Path(tempfile.mkdtemp(prefix="tailorresume_docx_pdf_"))
    try:
        working_docx = temp_dir / "styled_resume.docx"
        pdf_path = temp_dir / "styled_resume.pdf"
        shutil.copy2(source_docx, working_docx)
        apply_resume_to_docx(working_docx, resume)
        ok, message = export_pdf(working_docx, pdf_path, pdf_cfg or {})
        pdf_bytes = pdf_path.read_bytes() if ok and pdf_path.exists() else b""
        docx_bytes = working_docx.read_bytes()
        return {
            "pdf": pdf_bytes,
            "html": build_pdf_preview_html(pdf_bytes, message),
            "markdown": "",
            "docx": docx_bytes,
            "pdf_message": message,
        }
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def build_docx_template_pdf_bundle(profile: dict, output_dir: Path | str, pdf_cfg: dict | None = None) -> dict[str, bytes | str]:
    source_docx = _uploaded_resume_path(profile)
    if not source_docx:
        raise FileNotFoundError("no resume so must upload resume")
    temp_dir = Path(tempfile.mkdtemp(prefix="tailorresume_docx_template_pdf_"))
    try:
        working_docx = temp_dir / "resume_template.docx"
        pdf_path = temp_dir / "resume_template.pdf"
        shutil.copy2(source_docx, working_docx)
        ok, message = export_pdf(working_docx, pdf_path, pdf_cfg or {})
        pdf_bytes = pdf_path.read_bytes() if ok and pdf_path.exists() else b""
        docx_bytes = working_docx.read_bytes()
        return {
            "pdf": pdf_bytes,
            "html": build_pdf_preview_html(pdf_bytes, message),
            "markdown": "",
            "docx": docx_bytes,
            "pdf_message": message,
        }
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
